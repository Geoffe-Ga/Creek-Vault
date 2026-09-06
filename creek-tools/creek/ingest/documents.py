"""Document file ingestor for the Creek ingest pipeline.

Ingests DOCX, PDF, HTML, TXT, and RTF files by converting them to
markdown fragments. Uses format-specific libraries for extraction:

- **DOCX**: ``python-docx`` for paragraphs, tables, and headings
- **PDF**: ``pdfminer.six`` for text extraction; an image-only PDF is
  detected by ``_detect_scanned_pdf`` and routed to
  :meth:`creek.ingest.images.ImageIngestor.ingest_pdf` for per-page OCR
  (#1639), which is the only way its pages produce any body at all
- **HTML**: ``markdownify`` for direct HTML-to-markdown conversion
- **TXT**: Heuristic structure detection wrapping

Exports:
    DocumentIngestor: Concrete ``Ingestor`` subclass for document files.
    _parse_docx_to_markdown: Convert DOCX bytes to markdown string.
    _parse_pdf_to_text: Extract text from PDF bytes.
    _parse_html_to_markdown: Convert HTML string to markdown.
    _wrap_txt_with_structure: Wrap plain text with heuristic structure.
    _detect_scanned_pdf: Check if a PDF is likely scanned (image-only).
    _extract_docx_metadata: Extract metadata from DOCX bytes.
    _extract_pdf_metadata: Extract metadata from PDF bytes.
    _resolve_document_author: Split extracted author metadata into axis + name.
    _infer_document_platform: Map file extension to SourcePlatform.
"""

from __future__ import annotations

import contextlib
import io
import logging
import re
import string
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

from creek.ingest._detection import is_substack_export
from creek.ingest.base import (
    Ingestor,
    ParsedFragment,
    RawDocument,
    file_modified_time,
    normalize_encoding,
    normalize_timestamp,
    parse_authored_at,
)
from creek.ingest.html import extract_html_authored_at, parse_html_to_markdown
from creek.ingest.images import ImageIngestor
from creek.models import Authorship, SourcePlatform

if TYPE_CHECKING:
    from creek.config import OCRConfig

logger = logging.getLogger(__name__)

# ---- Constants ----

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {".docx", ".pdf", ".html", ".htm", ".txt", ".rtf"}
)
"""File extensions supported by the DocumentIngestor."""

_SCANNED_TEXT_THRESHOLD = 100
"""Minimum character count to consider a multi-page PDF as text-based."""

_DOCUMENT_PLATFORM_EXTENSIONS: frozenset[str] = frozenset({".docx", ".pdf", ".rtf"})
"""Extensions that map to the DOCUMENT source platform."""

_AUTHORSHIP_BY_VALUE: dict[str, Authorship] = {
    member.value: member for member in Authorship
}
"""Lookup from an ``Authorship`` member's wire value to the member itself."""


# ---- Helper Functions ----


def _parse_html_to_markdown(html: str) -> str:
    """Re-export of :func:`creek.ingest.html.parse_html_to_markdown`.

    The implementation moved to :mod:`creek.ingest.html` so
    :mod:`creek.ingest.substack` can use it without reaching into a
    private symbol of this module. This thin shim keeps existing
    tests and downstream callers that import the underscore-prefixed
    name working; new code should import the public
    :func:`creek.ingest.html.parse_html_to_markdown` directly.
    """
    return parse_html_to_markdown(html)


def _wrap_txt_with_structure(text: str) -> str:
    """Wrap plain text with heuristic structure detection.

    Detects potential titles (first non-empty line) and preserves
    list items. Returns the text with minimal markdown formatting.

    Args:
        text: The plain text content.

    Returns:
        Text with basic markdown structure, or empty string for blank input.
    """
    if not text.strip():
        return ""

    lines = text.strip().split("\n")
    result_lines: list[str] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        if i == 0 and stripped and not stripped.startswith("-"):
            # Treat first non-empty line as a heading
            result_lines.append(f"# {stripped}")
        else:
            result_lines.append(line)

    return "\n".join(result_lines)


def _extract_pdf_text_from_bytes(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfminer.

    Args:
        pdf_bytes: Raw PDF file bytes.

    Returns:
        Extracted text content.
    """
    from pdfminer.high_level import (
        extract_text,
    )

    result: str = extract_text(io.BytesIO(pdf_bytes))
    return result


def _count_pdf_pages(pdf_bytes: bytes) -> int:
    """Count the number of pages in a PDF.

    Args:
        pdf_bytes: Raw PDF file bytes.

    Returns:
        The number of pages in the PDF.
    """
    from pdfminer.pdfpage import (
        PDFPage,
    )

    return sum(1 for _ in PDFPage.get_pages(io.BytesIO(pdf_bytes)))


def _parse_pdf_to_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes.

    Delegates to ``_extract_pdf_text_from_bytes`` for the actual extraction.

    Args:
        pdf_bytes: Raw PDF file bytes.

    Returns:
        Extracted text content.
    """
    return _extract_pdf_text_from_bytes(pdf_bytes)


def _detect_scanned_pdf(text: str, page_count: int) -> bool:
    """Detect if a PDF is likely scanned (image-only).

    A PDF is flagged as scanned if it has multiple pages but the
    extracted text is shorter than the threshold (100 characters).

    Args:
        text: The extracted text from the PDF.
        page_count: The number of pages in the PDF.

    Returns:
        True if the PDF appears to be scanned.
    """
    return page_count > 1 and len(text.strip()) < _SCANNED_TEXT_THRESHOLD


def _parse_docx_to_markdown(docx_bytes: bytes) -> str:
    """Convert DOCX bytes to a markdown string.

    Iterates through paragraphs and tables in the document, converting
    headings to ATX-style markdown and tables to pipe-delimited format.

    Args:
        docx_bytes: Raw DOCX file bytes.

    Returns:
        A markdown-formatted string.
    """
    from docx import (
        Document as DocxDocument,
    )

    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            _convert_paragraph(element, doc, parts)
        elif tag == "tbl":
            _convert_table(element, doc, parts)

    return "\n\n".join(parts)


def _convert_paragraph(element: Any, doc: Any, parts: list[str]) -> None:
    """Convert a DOCX paragraph element to markdown and append to parts.

    Args:
        element: The paragraph XML element.
        doc: The python-docx Document object.
        parts: The list of markdown parts to append to.
    """
    from docx.text.paragraph import Paragraph

    para = Paragraph(element, doc)
    text = para.text.strip()
    if not text:
        return

    style_name = para.style.name if para.style else ""
    heading_level = _get_heading_level(style_name)

    if heading_level > 0:
        prefix = "#" * heading_level
        parts.append(f"{prefix} {text}")
    else:
        parts.append(text)


def _get_heading_level(style_name: str) -> int:
    """Extract heading level from a DOCX style name.

    Args:
        style_name: The paragraph style name (e.g., 'Heading 1').

    Returns:
        The heading level (1-6), or 0 if not a heading.
    """
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 0
    return 0


def _convert_table(element: Any, doc: Any, parts: list[str]) -> None:
    """Convert a DOCX table element to markdown and append to parts.

    Args:
        element: The table XML element.
        doc: The python-docx Document object.
        parts: The list of markdown parts to append to.
    """
    from docx.table import Table

    table = Table(element, doc)
    rows = table.rows
    if not rows:
        return

    table_lines: list[str] = []
    for i, row in enumerate(rows):
        cells = [cell.text.strip() for cell in row.cells]
        table_lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            table_lines.append("| " + " | ".join("---" for _ in cells) + " |")

    parts.append("\n".join(table_lines))


def _extract_docx_metadata(docx_bytes: bytes) -> dict[str, Any]:
    """Extract metadata from DOCX core properties.

    Extracts author, title, creation date, and modification date from
    the document's core properties. The modification date powers the
    FEAT-031 ``authored_at`` fallback when ``created`` is absent.

    Args:
        docx_bytes: Raw DOCX file bytes.

    Returns:
        A dict with author, title, created, and modified keys.
    """
    from docx import (
        Document as DocxDocument,
    )

    doc = DocxDocument(io.BytesIO(docx_bytes))
    props = doc.core_properties

    metadata: dict[str, Any] = {}
    if props.author:
        metadata["author"] = props.author
    if props.title:
        metadata["title"] = props.title
    if props.created:
        metadata["created_date"] = props.created.isoformat()
    if props.modified:
        metadata["modified_date"] = props.modified.isoformat()

    return metadata


def _parse_docx_authored_at(metadata: dict[str, Any]) -> datetime | None:
    """Resolve a DOCX file's ``authored_at`` from extracted metadata.

    Walks ``created_date`` (``dcterms:created``) then ``modified_date``
    (``dcterms:modified``); returns the first parseable candidate.
    Returns ``None`` when both keys are absent or unparseable — FEAT-031
    forbids guessing the date.
    """
    for key in ("created_date", "modified_date"):
        value = metadata.get(key)
        if not value:
            continue
        try:
            parsed = parse_authored_at(value)
        except ValueError:
            continue
        if parsed is not None:
            return parsed
    return None


def _extract_pdf_metadata_from_bytes(pdf_bytes: bytes) -> dict[str, Any]:
    """Extract metadata from PDF document info dictionary.

    Pulls ``Author``, ``Title``, ``CreationDate``, and ``ModDate``.
    ``ModDate`` powers the FEAT-031 ``authored_at`` fallback when
    ``CreationDate`` is absent.

    Args:
        pdf_bytes: Raw PDF file bytes.

    Returns:
        A dict with available metadata fields (lowercase keys).
    """
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser

    metadata: dict[str, Any] = {}
    try:
        parser = PDFParser(io.BytesIO(pdf_bytes))
        pdf_doc = PDFDocument(parser)
        for info in pdf_doc.info:
            if isinstance(info, dict):
                for key in ("Author", "Title", "CreationDate", "ModDate"):
                    value = info.get(key)
                    if value is not None:
                        decoded = (
                            value.decode("utf-8", errors="replace")
                            if isinstance(value, bytes)
                            else str(value)
                        )
                        metadata[key.lower()] = decoded
    except Exception:
        logger.debug("Could not extract PDF metadata")

    return metadata


# PDF dates are encoded as ``D:YYYYMMDDHHmmSS+HH'mm'`` per ISO 32000.
# We strip the literal ``D:`` prefix and the apostrophes in the
# timezone, then hand the result to ``datetime.strptime`` with two
# format candidates (with-tz and without). Returning ``None`` rather
# than guessing matches the FEAT-031 contract.
_PDF_DATE_FORMATS: tuple[tuple[str, int], ...] = (
    ("%Y%m%d%H%M%S%z", 14),
    ("%Y%m%d%H%M%S", 14),
    ("%Y%m%d%H%M", 12),
    ("%Y%m%d", 8),
)
"""Each unseparated format paired with the digit count it accepts.

The count is load-bearing, not documentation. ``datetime.strptime`` lets
every numeric directive match one *or* two digits, so an ungated
``%Y%m%d%H%M%S`` re-segments a malformed run into a plausible-looking but
**wrong** date rather than rejecting it: ``D:20241345`` became
2024-01-03 04:05 and ``D:99999999`` became 9999-09-09 09:09, both then
written to ``authored_at`` with no signal (#1632). A wrong date
propagates into threads, eddies and temporal linking, so guessing is the
one thing the FEAT-031 contract forbids.

The count covers only the leading digit run. The ``%z`` suffix is
variable-length (``+0500``, ``+05``, ``Z``), so gating on the total
string length would reject the canonical
``D:YYYYMMDDHHmmSS+HH'mm'`` form.
"""


def _parse_pdf_date(raw: str) -> datetime | None:
    """Parse a ``/CreationDate`` or ``/ModDate`` value into a tz-aware datetime.

    Accepts the canonical ``D:YYYYMMDDHHmmSS+HH'mm'`` format and a
    couple of common truncations. Returns ``None`` when no candidate
    matches — the caller falls through to the next field, never
    guesses.
    """
    cleaned = raw.strip().removeprefix("D:")
    # Strip the PDF spec's literal apostrophes from the tz offset:
    # ``+05'00'`` → ``+0500`` so ``%z`` parses it.
    cleaned = cleaned.replace("'", "")
    if cleaned.endswith(("z", "Z")):
        cleaned = cleaned[:-1] + "+0000"
    digits = len(cleaned) - len(cleaned.lstrip(string.digits))
    for fmt, expected_digits in _PDF_DATE_FORMATS:
        if digits != expected_digits:
            continue
        try:
            parsed = datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
        try:
            return parse_authored_at(parsed)
        except ValueError:
            return None
    return None


def _parse_pdf_authored_at(metadata: dict[str, Any]) -> datetime | None:
    """Resolve a PDF file's ``authored_at`` from extracted metadata.

    Walks ``creationdate`` then ``moddate`` (matching the per-format
    extraction chain in the FEAT-031 spec); each value is parsed via
    :func:`_parse_pdf_date`. Returns ``None`` when both are absent or
    unparseable.
    """
    for key in ("creationdate", "moddate"):
        raw = metadata.get(key)
        if not raw:
            continue
        parsed = _parse_pdf_date(str(raw))
        if parsed is not None:
            return parsed
    return None


# RTF date control words follow ``\creatim`` / ``\revtim`` and embed
# their fields via further control words (``\yr2024 \mo3 \dy15 \hr8
# \min30 \sec0``) before the closing ``}``. The regex captures the
# control word and the inner blob so the per-field parse can run on
# the captured group.
_RTF_DATE_GROUP = re.compile(
    r"\\(creatim|revtim)\b([^}]*)",
    re.IGNORECASE,
)
_RTF_DATE_FIELDS = (
    ("yr", "year"),
    ("mo", "month"),
    ("dy", "day"),
    ("hr", "hour"),
    ("min", "minute"),
    ("sec", "second"),
)


def _parse_rtf_authored_at(raw_bytes: bytes) -> datetime | None:
    """Extract ``\\creatim`` (then ``\\revtim``) from an RTF byte stream.

    RTF stores dates as a sequence of ``\\yr``/``\\mo``/``\\dy``/etc.
    control words inside ``\\creatim{…}``. Returns a UTC-anchored
    datetime when at least year + month + day were captured; falls
    through to ``\\revtim`` if ``\\creatim`` is absent or unparseable.
    The result is UTC because RTF carries no tz information — per the
    FEAT-031 spec, that is the honest default rather than guessing.
    """
    try:
        text = raw_bytes.decode("ascii", errors="replace")
    except UnicodeDecodeError:
        return None
    candidates: dict[str, dict[str, int]] = {}
    for match in _RTF_DATE_GROUP.finditer(text):
        word = match.group(1).lower()
        body = match.group(2)
        fields: dict[str, int] = {}
        for token, dest in _RTF_DATE_FIELDS:
            sub = re.search(rf"\\{token}(-?\d+)", body)
            if sub:
                with contextlib.suppress(ValueError):
                    fields[dest] = int(sub.group(1))
        candidates[word] = fields
    for word in ("creatim", "revtim"):
        fields = candidates.get(word, {})
        if not {"year", "month", "day"} <= fields.keys():
            continue
        try:
            naive = datetime(
                year=fields["year"],
                month=fields["month"],
                day=fields["day"],
                hour=fields.get("hour", 0),
                minute=fields.get("minute", 0),
                second=fields.get("second", 0),
            )
        except (TypeError, ValueError):
            continue
        try:
            return parse_authored_at(naive)
        except ValueError:
            continue
    return None


def _extract_pdf_metadata(pdf_bytes: bytes) -> dict[str, Any]:
    """Extract metadata from PDF bytes.

    Delegates to ``_extract_pdf_metadata_from_bytes``.

    Args:
        pdf_bytes: Raw PDF file bytes.

    Returns:
        A dict with available metadata fields.
    """
    return _extract_pdf_metadata_from_bytes(pdf_bytes)


def _resolve_document_author(
    raw_author: object,
) -> tuple[Authorship, str | None] | None:
    """Split a document's author metadata into an authorship axis and a name.

    ``FragmentSource.author`` answers "whose views does this stand for?" on the
    ``self|ai|other|collaborative`` axis. A DOCX ``core_properties.author`` or a
    PDF ``/Author`` answers the different question "what name is on the file".
    Copying the second into the first was issue #1229: every document Word saves
    carries an author, so every such document failed Pydantic validation and was
    dropped at assembly.

    Returns ``None`` when there is no usable value — absent, non-string, or
    blank — leaving the model's ``self`` default in force. "We know nothing" is
    not evidence about anybody.

    A value that already names an ``Authorship`` member is honoured as the axis
    and yields no name; it is a classification, not somebody's name. Any other
    string resolves the axis to :attr:`Authorship.OTHER` and is returned as the
    name. ``OTHER`` is both the convention this codebase already applies to an
    explicit author name
    (:meth:`creek.clean.authorship.AuthorshipTagger._tag_self_platform`) and the
    fail-closed choice: ``SELF`` is what unlocks the INTIMATE privacy tier
    (:mod:`creek.classify.privacy`) and voice/skill generation
    (:mod:`creek.generate.skills`), so guessing ``SELF`` for a stranger's
    document would feed both from material that is not the owner's, whereas
    guessing ``OTHER`` for the owner's own document merely under-uses it — and
    the returned name is what makes that recoverable.

    Args:
        raw_author: The ``author`` value an ``_extract_*`` pass put in metadata.

    Returns:
        ``(authorship, name)`` where ``name`` is ``None`` for a value that was
        already an axis member, or ``None`` when there is nothing to record.
    """
    if not isinstance(raw_author, str):
        return None
    name = raw_author.strip()
    if not name:
        return None
    axis = _AUTHORSHIP_BY_VALUE.get(name.lower())
    if axis is not None:
        return axis, None
    return Authorship.OTHER, name


def _infer_document_platform(extension: str) -> SourcePlatform:
    """Map a file extension to a SourcePlatform value.

    DOCX, PDF, and RTF map to ``DOCUMENT``. HTML and TXT map to ``OTHER``.

    Args:
        extension: The file extension (e.g., '.docx').

    Returns:
        The corresponding SourcePlatform enum value.
    """
    if extension in _DOCUMENT_PLATFORM_EXTENSIONS:
        return SourcePlatform.DOCUMENT
    return SourcePlatform.OTHER


# ---- DocumentIngestor ----


class DocumentIngestor(Ingestor):
    """Ingestor for common document formats (DOCX, PDF, HTML, TXT, RTF).

    Discovers files with supported extensions, parses them using
    format-specific libraries, converts to markdown, and generates
    Creek-compatible frontmatter with metadata extraction.
    """

    def __init__(self, *, ocr: OCRConfig | None = None) -> None:
        """Wire the scanned-PDF OCR route, or deliberately decline to.

        **Fail-closed is the shape of this constructor, not a check inside
        it.** ``ocr is None`` means "the caller has no vault ``ocr`` block",
        which is what every MCP surface passes today
        (``creek_mcp.tools.upload`` and its siblings supply none). Mirroring
        :func:`~creek.ingest.pipeline.build_ingestor`'s image default — where
        ``None`` means "use OCR defaults", i.e. on — would let a remote caller
        obtain OCR in a vault whose operator had written ``ocr.enabled:
        false``. So this route is off unless a block explicitly turns it on,
        which is deliberately stricter than the image path and preserves
        today's behaviour byte-for-byte for every caller that passes nothing.

        Resolution is **eager**, so an unknown ``ocr.engine`` or an empty
        ``ocr.languages`` raises
        :class:`~creek.ingest.images.OcrConfigError` out of *construction* —
        outside :meth:`Ingestor._parse_safe`, which would otherwise swallow it
        into ``IngestResult.errors`` and let the run report a success it did
        not have. It reaches the CLI's exit 2 instead. Construction is cheap:
        :class:`~creek.ingest.images.PytesseractOcrEngine` defers every
        import.

        ``enabled: false`` resolves no engine at all, for the reason
        :meth:`~creek.ingest.images.ImageIngestor.from_ocr_config` records: an
        operator who has switched OCR off must not be refused for the name of
        a backend they just said they did not want.

        Args:
            ocr: The vault's ``ocr`` block, when the caller has one.

        Raises:
            UnknownOcrEngineError: When ``ocr.engine`` names no known backend
                and ``ocr.enabled`` is true.
            OcrConfigError: When ``ocr.languages`` holds no usable code.
        """
        self._pdf_ocr: ImageIngestor | None = (
            ImageIngestor.from_ocr_config(ocr)
            if ocr is not None and ocr.enabled
            else None
        )
        self._pdf_ocr_declined: bool = ocr is not None and not ocr.enabled

    def discover(self, source_path: Path) -> list[RawDocument]:
        """Find all document files at the given source path (recursively).

        If ``source_path`` is a file with a supported extension, returns
        a single-element list. If it is a directory, recursively searches
        for files with supported extensions. Returns empty list for
        nonexistent paths.

        Substack export directories (``posts.csv`` + per-post
        ``<id>.<slug>.html``) are claimed by ``SubstackIngestor``;
        ``DocumentIngestor`` defers to it rather than emitting every
        post a second time as an opaque HTML document. Since issue #1304
        the pipeline would arbitrate that contest anyway — ``substack``
        outranks ``document`` in
        :data:`creek.ingest.routing.CLAIM_PRIORITY` — but narrowing the
        claim here is still worth doing: it saves parsing a whole export
        directory whose output is guaranteed to be discarded, and it
        keeps ``creek ingest --type document`` honest, which the
        arbiter never sees.

        Args:
            source_path: A file or directory path to search.

        Returns:
            A list of ``RawDocument`` objects for each discovered file.
        """
        if not source_path.exists():
            return []

        if source_path.is_file():
            return self._discover_single_file(source_path)

        if is_substack_export(source_path):
            return []

        return self._discover_directory(source_path)

    def _discover_single_file(self, file_path: Path) -> list[RawDocument]:
        """Discover a single document file.

        Args:
            file_path: Path to the file.

        Returns:
            A single-element list if the file has a supported extension.
        """
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return []

        return [self._read_file(file_path)]

    def _discover_directory(self, dir_path: Path) -> list[RawDocument]:
        """Recursively discover all document files in a directory.

        Args:
            dir_path: Directory path to search.

        Returns:
            A list of RawDocument objects for each document file found.
        """
        docs: list[RawDocument] = [
            self._read_file(file_path)
            for file_path in sorted(dir_path.rglob("*"))
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
        ]
        return docs

    def _read_file(self, file_path: Path) -> RawDocument:
        """Read a file into a RawDocument.

        Args:
            file_path: Path to the file.

        Returns:
            A RawDocument instance.
        """
        raw_bytes = file_path.read_bytes()
        _text, encoding = normalize_encoding(raw_bytes)
        return RawDocument(
            path=file_path,
            content=raw_bytes,
            metadata={
                "source_type": "document",
                "file_type": file_path.suffix.lower(),
            },
            detected_encoding=encoding,
        )

    def parse(self, raw: RawDocument) -> list[ParsedFragment]:
        """Parse a raw document, routing to format-specific handlers.

        Dispatches to DOCX, PDF, HTML, or TXT parsers based on the
        file extension. Extracts metadata where available.

        FEAT-031: per-format ``authored_at`` extraction chains:

        * **HTML / HTM**: ``<meta property="article:published_time">``,
          OpenGraph, Dublin Core, ``<meta name="date">``, JSON-LD
          ``datePublished``.
        * **PDF**: ``/CreationDate`` (then ``/ModDate``) from the
          document info dictionary.
        * **DOCX**: core properties ``dcterms:created`` (then
          ``dcterms:modified``).
        * **RTF**: ``\\creatim`` (then ``\\revtim``).
        * **TXT**: filesystem mtime only — no embedded date metadata
          exists, so ``authored_at`` is ``None`` and downstream
          surfaces fall through to ``ingested``.

        #1639: an image-only PDF returns **N** fragments, one per page,
        not one. Metadata is extracted first so ``scanned`` is known before
        the decision, and the scanned branch returns before
        :meth:`_extract_content` runs — which is also what keeps the route to
        exactly one ``_parse_pdf_to_text`` call rather than the two an
        unrouted PDF pays (``_extract_pdf_content`` and ``_add_pdf_metadata``
        each run it). Reordering is safe: :meth:`_extract_content` reads only
        ``file_type``, ``raw_bytes`` and ``text``, and none of the metadata
        helpers reads content.

        Args:
            raw: The raw document to parse.

        Returns:
            A single-element list for every format but a scanned PDF, which
            returns one fragment per non-empty page.
        """
        file_type = raw.metadata.get("file_type", raw.path.suffix.lower())
        text, encoding = normalize_encoding(raw.content)

        metadata = self._extract_metadata(file_type, raw.content, encoding)
        metadata["authored_at"] = self._extract_authored_at(
            file_type, raw.content, text, metadata
        )
        timestamp = self._resolve_timestamp(metadata, raw.path)

        # A PDF whose page count could not be read never gains ``scanned``
        # (``_add_pdf_metadata`` swallows that), so it falls through here and
        # ingests exactly as it does today.
        if metadata.get("scanned"):
            if self._pdf_ocr is not None:
                pages = self._ocr_scanned_pdf(
                    self._pdf_ocr, raw.path, metadata, timestamp
                )
                # ``ingest_pdf`` skips pages OCR recovered nothing from, so a
                # scan it could read nothing at all from comes back empty.
                # Falling through rather than returning ``[]`` is deliberate:
                # the operator then gets the same single ``scanned: true``
                # fragment an un-routed scan produces, instead of the file
                # disappearing from the run without a word.
                if pages:
                    return pages
            elif self._pdf_ocr_declined:
                metadata["ocr_declined"] = True

        content = self._extract_content(file_type, raw.content, text)

        return [
            ParsedFragment(
                content=content,
                metadata=metadata,
                source_path=str(raw.path),
                timestamp=timestamp,
            )
        ]

    def _ocr_scanned_pdf(
        self,
        ocr: ImageIngestor,
        pdf_path: Path,
        metadata: dict[str, Any],
        timestamp: datetime,
    ) -> list[ParsedFragment]:
        """OCR an image-only PDF into one fragment per page (#1639).

        The per-page fragments come back from
        :meth:`~creek.ingest.images.ImageIngestor.ingest_pdf` and are
        **mutated in place rather than rebuilt**, deliberately: rebuilding
        would drop the ``source_unit`` that method mints, and that field is
        the whole of this fragment's identity under the ``document`` ledger.

        The two metadata dicts are merged with the OCR side winning, so the
        page keeps ``page`` / ``ocr_confidence`` / ``review`` while gaining
        the document keys this ingestor's own renderers read — ``file_type``
        (which decides the platform), ``source_encoding``, ``scanned`` and
        ``authored_at``. The title is set last and per page: a scanned PDF
        with a ``/Title`` would otherwise give every page the same one.
        That is a **legibility** requirement, not a loss-prevention one —
        measured, the writer disambiguates identical titles to ``-1`` /
        ``-2`` suffixes and loses no content.

        The timestamp is the document route's, not the OCR route's, so a
        page's id is anchored to the same instant an unrouted PDF's would be.

        Args:
            ocr: The image ingestor resolved from the vault's ``ocr`` block.
            pdf_path: The scanned PDF.
            metadata: The document metadata extracted for the whole file.
            timestamp: The identity anchor resolved for the whole file.

        Returns:
            One fragment per non-empty page.
        """
        stem = pdf_path.stem
        fragments = ocr.ingest_pdf(pdf_path)
        for fragment in fragments:
            page = fragment.metadata.get("page")
            fragment.metadata = (
                metadata | fragment.metadata | {"title": f"{stem} — page {page}"}
            )
            fragment.timestamp = timestamp
        return fragments

    def _extract_authored_at(
        self,
        file_type: str,
        raw_bytes: bytes,
        text: str,
        metadata: dict[str, Any],
    ) -> datetime | None:
        """Route to per-format authored-date extraction (FEAT-031).

        Returns ``None`` when no source date exists or the candidate
        parsers fail. The TXT branch is the canonical "no embedded
        metadata" case; HTML / PDF / DOCX / RTF each consult their
        native metadata before falling through.
        """
        if file_type in {".html", ".htm"}:
            raw_value = extract_html_authored_at(text)
            if raw_value:
                try:
                    return parse_authored_at(raw_value)
                except ValueError:
                    return None
            return None
        if file_type == ".pdf":
            return _parse_pdf_authored_at(metadata)
        if file_type == ".docx":
            return _parse_docx_authored_at(metadata)
        if file_type == ".rtf":
            return _parse_rtf_authored_at(raw_bytes)
        return None

    def _extract_content(self, file_type: str, raw_bytes: bytes, text: str) -> str:
        """Extract content based on file type.

        Args:
            file_type: The file extension.
            raw_bytes: Raw file bytes.
            text: Decoded text content.

        Returns:
            Extracted content string.
        """
        if file_type == ".docx":
            return _parse_docx_to_markdown(raw_bytes)
        if file_type == ".pdf":
            return self._extract_pdf_content(raw_bytes)
        if file_type in {".html", ".htm"}:
            return parse_html_to_markdown(text)
        if file_type == ".txt":
            return _wrap_txt_with_structure(text)
        # RTF and other formats: return raw text
        return text

    def _extract_pdf_content(self, raw_bytes: bytes) -> str:
        """Extract text content from PDF bytes.

        One delegation and nothing else. Scanned *detection* lives in
        :meth:`_add_pdf_metadata` and the OCR route it feeds is taken in
        :meth:`parse` before this method is reached, so a PDF that arrives
        here is one whose text pdfminer can actually read.

        Args:
            raw_bytes: Raw PDF file bytes.

        Returns:
            Extracted text content.
        """
        return _parse_pdf_to_text(raw_bytes)

    def _extract_metadata(
        self, file_type: str, raw_bytes: bytes, encoding: str
    ) -> dict[str, Any]:
        """Extract metadata based on file type.

        Args:
            file_type: The file extension.
            raw_bytes: Raw file bytes.
            encoding: Detected character encoding.

        Returns:
            Metadata dict with file_type, source_encoding, and format-specific fields.
        """
        metadata: dict[str, Any] = {
            "file_type": file_type,
            "source_encoding": encoding,
        }

        if file_type == ".docx":
            metadata.update(_extract_docx_metadata(raw_bytes))
        elif file_type == ".pdf":
            self._add_pdf_metadata(raw_bytes, metadata)

        return metadata

    def _add_pdf_metadata(self, raw_bytes: bytes, metadata: dict[str, Any]) -> None:
        """Add PDF-specific metadata including scanned detection.

        Args:
            raw_bytes: Raw PDF file bytes.
            metadata: Metadata dict to update.
        """
        pdf_meta = _extract_pdf_metadata(raw_bytes)
        metadata.update(pdf_meta)

        try:
            text = _parse_pdf_to_text(raw_bytes)
            page_count = _count_pdf_pages(raw_bytes)
            if _detect_scanned_pdf(text, page_count):
                metadata["scanned"] = True
        except Exception:
            logger.debug("Could not check for scanned PDF")

    def _resolve_timestamp(self, metadata: dict[str, Any], file_path: Path) -> datetime:
        """Resolve a timestamp from metadata or filesystem.

        Checks metadata for created_date first, then falls back to
        the file's modification time.

        **This value is an identity anchor, not an authorship claim.**
        :func:`creek.ingest.base.generate_fragment_id` hashes it, so the
        fallback must be a pure function of the file's epoch mtime —
        invariant under the host's ``TZ`` env var, its tzdata, its DST
        state, and its operating system (#1329). That is exactly what
        :func:`~creek.ingest.base.file_modified_time` guarantees, and it is
        why this must never be "simplified" back to a bare
        ``datetime.fromtimestamp(mtime)``, which renders the epoch in the
        host's local zone and so mints one id per timezone.

        Documents used to be more exposed to this than markdown, not less:
        :func:`creek.ingest.pipeline.ledger_for_source` left them
        unledgered, so there was no recorded id to fall back on and a
        drifting derivation was an unconditional second write rather than a
        merely possible one. #1363 closed that — ``document`` is now in
        :data:`creek.ingest.pipeline.LEDGERED_SOURCES`, once #953 had made
        ``derive_source_key`` collision-free for out-of-vault sources, which
        was the blocker. A ledgered document's id comes from the record, so
        the derivation below is only how a *new* document gets its first
        one; keeping it host-independent still matters for exactly that.

        Args:
            metadata: Fragment metadata dict.
            file_path: Path to the source file.

        Returns:
            A timezone-aware datetime; UTC when it comes from the filesystem.
        """
        created = metadata.get("created_date")
        if created is not None:
            try:
                return normalize_timestamp(str(created), None)
            except ValueError:
                logger.warning("Invalid metadata timestamp: %s", created)

        return file_modified_time(file_path)

    def convert_to_markdown(self, fragment: ParsedFragment) -> str:
        """Return the fragment content as markdown.

        Content is already converted to markdown during parsing, so this
        method returns it as-is — except for a scanned-PDF page, which is
        handed back to :meth:`creek.ingest.images.ImageIngestor.convert_to_markdown`
        so it renders with the Obsidian ``![[file#page=N]]`` embed that
        method exists to produce. Without the delegation that branch stays
        dead production code after the routing lands, and no gate would say
        so: it is a branch, not a symbol, so vulture is silent about it.

        Keyed on the ``page`` metadata only an OCR'd PDF page carries, so
        every other document still returns its content untouched.

        Args:
            fragment: The parsed fragment to convert.

        Returns:
            The markdown content string.
        """
        if self._pdf_ocr is not None and fragment.metadata.get("page"):
            return self._pdf_ocr.convert_to_markdown(fragment)
        return fragment.content

    def generate_frontmatter(self, fragment: ParsedFragment) -> dict[str, Any]:
        """Generate Creek-compatible YAML frontmatter for a document fragment.

        Builds frontmatter with type, title, source (platform, original_file,
        original_encoding), created timestamp, and — for a scanned PDF — the
        top-level ``scanned``, ``page`` and ``review`` markers.

        An extracted document author is split by :func:`_resolve_document_author`
        into ``source.author`` (the ``Authorship`` axis) and ``source.author_name``
        (the free-text name), never conflated into one slot (#1229).

        **``scanned`` is emitted top-level, and that is a fix, not a style
        choice (#1639).** It used to be written as ``source["scanned"]``, and
        measured at HEAD it never reached a single vault file: ``FragmentSource``
        does not model it, ``Fragment.model_validate`` leaves pydantic's
        ``extra="ignore"`` in place, and
        :data:`~creek.ingest.base.PASSTHROUGH_FRONTMATTER_KEYS` is a *top-level*
        allowlist that cannot rescue a key nested under ``source``. So an
        operator whose PDF was an image got an empty fragment and no indication
        of any kind. The three keys below ride that allowlist by presence, so a
        document with nothing to say about them gains no line at all.

        Args:
            fragment: The parsed fragment with metadata.

        Returns:
            A dict of frontmatter key-value pairs.
        """
        file_type = fragment.metadata.get("file_type", "")
        platform = _infer_document_platform(file_type)
        title = fragment.metadata.get("title", Path(fragment.source_path).stem)

        source: dict[str, Any] = {
            "platform": platform,
            "original_file": fragment.source_path,
            "original_encoding": fragment.metadata.get("source_encoding", "utf-8"),
        }

        # Add optional metadata fields
        resolved_author = _resolve_document_author(fragment.metadata.get("author"))
        if resolved_author is not None:
            authorship, author_name = resolved_author
            source["author"] = authorship
            if author_name is not None:
                source["author_name"] = author_name
        frontmatter_dict: dict[str, Any] = {
            "type": "fragment",
            "title": title,
            "source": source,
            "created": fragment.timestamp.isoformat(),
        }
        if fragment.metadata.get("scanned"):
            frontmatter_dict["scanned"] = True
        page = fragment.metadata.get("page")
        if page:
            frontmatter_dict["page"] = page
        review = fragment.metadata.get("review")
        if review:
            frontmatter_dict["review"] = review
        authored_at: datetime | None = fragment.metadata.get("authored_at")
        if authored_at is not None:
            frontmatter_dict["authored_at"] = authored_at.isoformat()
        return frontmatter_dict
