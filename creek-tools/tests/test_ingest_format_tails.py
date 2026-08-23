"""Structure-conversion tails for the document/deck/sheet ingestors.

These are the non-date branches of the same three files #855, #856 and
#870 were opened against — empty paragraphs, unnumbered heading styles,
header-less sheets, slides with no body. They are what gives those
files a *durable* margin over the 80% strict per-file gate rather than
an accidental one, which is the precondition for retiring their entries
from ``scripts/coverage-waivers.txt``.

The date chain itself lives in ``test_ingest_authored_at_chain.py``;
this module deliberately holds no date assertions.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import TYPE_CHECKING, Any

import pytest

from creek.ingest.base import ParsedFragment
from creek.ingest.documents import (
    DocumentIngestor,
    _extract_docx_metadata,
    _extract_pdf_metadata_from_bytes,
    _get_heading_level,
    _parse_docx_authored_at,
    _parse_docx_to_markdown,
)
from creek.ingest.presentations import (
    PresentationData,
    PresentationIngestor,
    SlideData,
    _slide_to_data,
)
from creek.ingest.spreadsheets import (
    SheetData,
    SpreadsheetIngestor,
    _cell_to_str,
    _split_header,
)

if TYPE_CHECKING:
    from pathlib import Path


# ---- DOCX markdown structure -------------------------------------------


class TestDocxMarkdownStructure:
    """Empty and odd elements are skipped, not rendered as blanks."""

    def test_blank_paragraphs_are_dropped(self, tmp_path: Path) -> None:
        """A whitespace-only paragraph contributes nothing to the markdown."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Real content.")
        doc.add_paragraph("   ")
        doc.add_paragraph("")
        doc.add_paragraph("More content.")
        path = tmp_path / "blanks.docx"
        doc.save(str(path))

        result = _parse_docx_to_markdown(path.read_bytes())

        assert "Real content." in result
        assert "More content." in result
        assert "\n\n\n" not in result

    def test_empty_table_is_dropped(self, tmp_path: Path) -> None:
        """A zero-row table emits no separator line."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Before.")
        doc.add_table(rows=0, cols=3)
        doc.add_paragraph("After.")
        path = tmp_path / "emptytable.docx"
        doc.save(str(path))

        result = _parse_docx_to_markdown(path.read_bytes())

        assert "Before." in result
        assert "After." in result
        assert "| --- |" not in result

    @pytest.mark.parametrize(
        ("style_name", "expected"),
        [
            pytest.param("Heading 1", 1, id="heading-1"),
            pytest.param("Heading 6", 6, id="heading-6"),
            pytest.param("Normal", 0, id="not-a-heading"),
            pytest.param("", 0, id="empty-style"),
            pytest.param("Heading", 0, id="heading-with-no-number"),
            pytest.param("Heading Foo", 0, id="heading-with-non-numeric-suffix"),
        ],
    )
    def test_heading_level_extraction(self, style_name: str, expected: int) -> None:
        """A style name that is not ``Heading <n>`` degrades to level 0.

        ``"Heading"`` alone hits the ``IndexError``-adjacent path
        (``split()[-1]`` yields ``"Heading"``, which ``int()`` rejects)
        and ``"Heading Foo"`` hits the ``ValueError`` — both must return
        0 rather than propagate out of the markdown conversion.
        """
        assert _get_heading_level(style_name) == expected


def _bare_docx() -> Any:
    """Return a new DOCX whose core properties are genuinely empty.

    A fresh ``python-docx`` document ships with ``author="python-docx"``
    and both ``dcterms`` dates pre-populated, and the library refuses a
    ``None`` assignment to either date. Dropping the XML elements is the
    only way to reach the absent-property arms of
    :func:`_extract_docx_metadata`.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    element = doc.core_properties._element
    for child in list(element):
        if child.tag.endswith(("}created", "}modified")):
            element.remove(child)
    doc.core_properties.author = ""
    doc.core_properties.title = ""
    return doc


class TestDocxCoreProperties:
    """Each core property is emitted only when actually populated."""

    def test_absent_properties_are_omitted(self, tmp_path: Path) -> None:
        """A doc with no author/title/dates emits none of those keys."""
        doc = _bare_docx()
        path = tmp_path / "bare.docx"
        doc.save(str(path))

        metadata = _extract_docx_metadata(path.read_bytes())

        assert "author" not in metadata
        assert "title" not in metadata
        assert "created_date" not in metadata
        assert "modified_date" not in metadata

    def test_author_without_title(self, tmp_path: Path) -> None:
        """Author present, title absent — only ``author`` appears."""
        doc = _bare_docx()
        doc.core_properties.author = "Ada"
        path = tmp_path / "author.docx"
        doc.save(str(path))

        metadata = _extract_docx_metadata(path.read_bytes())

        assert metadata["author"] == "Ada"
        assert "title" not in metadata

    def test_created_without_modified(self, tmp_path: Path) -> None:
        """``created`` present, ``modified`` absent — only the one key."""
        doc = _bare_docx()
        doc.core_properties.created = datetime(2024, 3, 15, tzinfo=UTC)
        path = tmp_path / "created.docx"
        doc.save(str(path))

        metadata = _extract_docx_metadata(path.read_bytes())

        assert metadata["created_date"].startswith("2024-03-15")
        assert "modified_date" not in metadata

    def test_whitespace_created_date_falls_through(self) -> None:
        """A truthy-but-blank ``created_date`` parses to ``None``, not a date.

        ``"   "`` passes the falsy guard but ``parse_authored_at``
        returns ``None`` for it, so the chain must keep walking to
        ``modified_date`` instead of returning ``None`` outright.
        """
        result = _parse_docx_authored_at(
            {"created_date": "   ", "modified_date": "2020-01-01T00:00:00+00:00"}
        )
        assert result == datetime(2020, 1, 1, tzinfo=UTC)


class TestPdfMetadataExtraction:
    """The PDF info dictionary is defensive about shape and content."""

    def test_junk_bytes_yield_empty_metadata(self) -> None:
        """pdfminer raises a wide error set on junk; all of it is swallowed."""
        assert _extract_pdf_metadata_from_bytes(b"definitely not a pdf") == {}

    def test_empty_bytes_yield_empty_metadata(self) -> None:
        """Zero-length input is a parse failure, not a crash."""
        assert _extract_pdf_metadata_from_bytes(b"") == {}

    def test_non_dict_info_entries_are_skipped(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """``pdf_doc.info`` is a list; non-dict entries carry nothing.

        ``_extract_pdf_metadata_from_bytes`` imports pdfminer *inside*
        the function, so patching the class on its home module is
        enough — the real production body still runs.
        """
        _patch_pdf_document(monkeypatch, info=["a string", 42, {"Title": "Real Title"}])
        assert _extract_pdf_metadata_from_bytes(b"x") == {"title": "Real Title"}

    def test_bytes_values_are_decoded(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """A ``bytes`` info value is decoded, not stringified as ``b'..'``."""
        _patch_pdf_document(
            monkeypatch, info=[{"Author": b"Ada Lovelace", "ModDate": None}]
        )
        assert _extract_pdf_metadata_from_bytes(b"x") == {"author": "Ada Lovelace"}

    def test_undecodable_bytes_are_replaced_not_raised(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Invalid UTF-8 in the info dict is replaced, never fatal."""
        _patch_pdf_document(monkeypatch, info=[{"Title": b"caf\xff"}])
        result = _extract_pdf_metadata_from_bytes(b"x")
        assert result["title"].startswith("caf")

    def test_non_bytes_values_are_stringified(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A non-bytes info value goes through ``str()``."""
        _patch_pdf_document(monkeypatch, info=[{"Title": 1234}])
        assert _extract_pdf_metadata_from_bytes(b"x") == {"title": "1234"}


def _patch_pdf_document(monkeypatch: pytest.MonkeyPatch, *, info: list[object]) -> None:
    """Point pdfminer's ``PDFDocument`` at a stub exposing *info*."""
    import pdfminer.pdfdocument
    import pdfminer.pdfparser

    class _FakeDoc:
        """PDFDocument stand-in with a caller-supplied ``info`` list."""

        def __init__(self, _parser: object) -> None:
            """Accept and ignore the parser, as PDFDocument does."""
            self.info = info

    class _FakeParser:
        """PDFParser stand-in that accepts any stream."""

        def __init__(self, _stream: object) -> None:
            """Accept and ignore the byte stream."""

    monkeypatch.setattr(pdfminer.pdfdocument, "PDFDocument", _FakeDoc)
    monkeypatch.setattr(pdfminer.pdfparser, "PDFParser", _FakeParser)


class TestResolveTimestamp:
    """An unusable ``created_date`` falls back to the file's mtime."""

    def test_invalid_created_date_falls_back_to_mtime(self, tmp_path: Path) -> None:
        """A junk metadata timestamp warns and falls back, never raises.

        The fallback must stay a pure function of the epoch mtime — it
        is hashed into the fragment id (#1329).
        """
        path = tmp_path / "note.txt"
        path.write_text("body", encoding="utf-8")
        ingestor = DocumentIngestor()

        resolved = ingestor._resolve_timestamp({"created_date": "not-a-date"}, path)

        assert resolved.tzinfo is not None
        assert resolved == ingestor._resolve_timestamp({}, path)

    def test_valid_created_date_is_used(self, tmp_path: Path) -> None:
        """The positive arm, so the fallback is not masking a total failure."""
        path = tmp_path / "note.txt"
        path.write_text("body", encoding="utf-8")
        ingestor = DocumentIngestor()

        resolved = ingestor._resolve_timestamp(
            {"created_date": "2024-03-15T08:30:00+00:00"}, path
        )

        assert resolved != ingestor._resolve_timestamp({}, path)


class TestScannedPdfDetection:
    """Scanned-PDF flagging is best-effort and never fatal."""

    def test_text_free_pdf_is_flagged_scanned(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A page-bearing PDF with no extractable text is marked scanned."""
        import creek.ingest.documents as documents_mod

        monkeypatch.setattr(documents_mod, "_extract_pdf_metadata", lambda _b: {})
        monkeypatch.setattr(documents_mod, "_parse_pdf_to_text", lambda _b: "")
        monkeypatch.setattr(documents_mod, "_count_pdf_pages", lambda _b: 5)

        metadata: dict[str, Any] = {}
        DocumentIngestor()._add_pdf_metadata(b"x", metadata)

        assert metadata["scanned"] is True

    def test_text_bearing_pdf_is_not_flagged(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A normal PDF carries no ``scanned`` key at all."""
        import creek.ingest.documents as documents_mod

        monkeypatch.setattr(documents_mod, "_extract_pdf_metadata", lambda _b: {})
        monkeypatch.setattr(
            documents_mod, "_parse_pdf_to_text", lambda _b: "lots of real text " * 50
        )
        monkeypatch.setattr(documents_mod, "_count_pdf_pages", lambda _b: 1)

        metadata: dict[str, Any] = {}
        DocumentIngestor()._add_pdf_metadata(b"x", metadata)

        assert "scanned" not in metadata

    def test_detection_failure_is_swallowed(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A raising page count loses the flag, not the document."""
        import creek.ingest.documents as documents_mod

        def _boom(_b: bytes) -> int:
            raise RuntimeError("pdfminer exploded")

        monkeypatch.setattr(documents_mod, "_extract_pdf_metadata", lambda _b: {})
        monkeypatch.setattr(documents_mod, "_parse_pdf_to_text", lambda _b: "")
        monkeypatch.setattr(documents_mod, "_count_pdf_pages", _boom)

        metadata: dict[str, Any] = {"file_type": ".pdf"}
        DocumentIngestor()._add_pdf_metadata(b"x", metadata)

        assert "scanned" not in metadata
        assert metadata["file_type"] == ".pdf"


# ---- Slide decomposition ------------------------------------------------


class _StubTextFrame:
    """Text-frame stand-in exposing a plain ``text`` attribute."""

    def __init__(self, text: str) -> None:
        """Store the frame's text verbatim."""
        self.text = text


class _StubShape:
    """Shape stand-in with a controllable text frame."""

    def __init__(self, text: str | None) -> None:
        """A ``None`` *text* means the shape has no text frame at all."""
        self.has_text_frame = text is not None
        if text is not None:
            self.text_frame = _StubTextFrame(text)


class _AttributelessShape:
    """Shape stand-in that omits ``has_text_frame`` altogether.

    python-pptx connectors and group shapes can lack the attribute
    rather than set it to ``False``, so the ``getattr`` default in
    :func:`_slide_to_data` is a live path and not decoration. Without
    this stub the default is never exercised — every other stub here
    sets the attribute explicitly.
    """


class _StubShapes(list[object]):
    """Shape collection stand-in that also exposes ``.title``."""

    def __init__(self, shapes: list[object], title: object | None) -> None:
        """Seed the collection and designate its title placeholder."""
        super().__init__(shapes)
        self.title = title


class _StubSlide:
    """Slide stand-in for :func:`_slide_to_data`."""

    def __init__(
        self,
        shapes: _StubShapes,
        *,
        notes: str | None = None,
        notes_raise: bool = False,
    ) -> None:
        """Wire up the shapes and the optional notes slide."""
        self.shapes = shapes
        self.has_notes_slide = notes is not None or notes_raise
        if notes_raise:
            self.notes_slide = object()
        elif notes is not None:
            self.notes_slide = _StubNotes(notes)


class _BareSlide:
    """Slide stand-in that omits ``has_notes_slide`` but *has* notes.

    The companion to :class:`_AttributelessShape`: it is the only way
    to reach the ``getattr(slide, "has_notes_slide", False)`` default,
    which :class:`_StubSlide` always sets explicitly. It deliberately
    still carries a populated ``notes_slide`` so the default is
    *observable*: a slide that never declared it has notes must not
    have its notes read, and flipping the default to ``True`` changes
    the answer instead of being swallowed by the ``AttributeError``
    arm.
    """

    def __init__(self, shapes: _StubShapes) -> None:
        """Expose the shapes and an unadvertised notes slide."""
        self.shapes = shapes
        self.notes_slide = _StubNotes("Notes nobody advertised")


class _StubNotes:
    """Notes-slide stand-in exposing ``notes_text_frame``."""

    def __init__(self, text: str) -> None:
        """Store the notes text."""
        self.notes_text_frame = _StubTextFrame(text)


class TestSlideDecomposition:
    """Shapes without text, and notes that misbehave, degrade cleanly."""

    def test_shape_without_text_frame_is_skipped(self) -> None:
        """A picture or line shape contributes nothing to the body."""
        body = _StubShape("Body text")
        slide = _StubSlide(_StubShapes([_StubShape(None), body], title=None))

        data = _slide_to_data(slide, 1)

        assert data.body == "Body text"
        assert data.title is None

    def test_shape_missing_has_text_frame_attribute_is_skipped(self) -> None:
        """A shape lacking the attribute hits the ``getattr`` default.

        Every other stub sets ``has_text_frame`` explicitly, so without
        this case flipping the default to ``True`` is a mutation no
        test can detect.
        """
        body = _StubShape("Body text")
        slide = _StubSlide(_StubShapes([_AttributelessShape(), body], title=None))

        assert _slide_to_data(slide, 1).body == "Body text"

    def test_blank_shape_text_is_skipped(self) -> None:
        """A whitespace-only text frame is not appended as an empty chunk."""
        slide = _StubSlide(
            _StubShapes([_StubShape("   "), _StubShape("Real")], title=None)
        )

        assert _slide_to_data(slide, 1).body == "Real"

    def test_title_placeholder_becomes_the_title(self) -> None:
        """The designated title shape populates ``title``, not ``body``."""
        title_shape = _StubShape("Slide Title")
        body_shape = _StubShape("Slide body")
        slide = _StubSlide(_StubShapes([title_shape, body_shape], title=title_shape))

        data = _slide_to_data(slide, 3)

        assert data.title == "Slide Title"
        assert data.body == "Slide body"
        assert data.index == 3

    def test_second_title_shaped_text_goes_to_body(self) -> None:
        """Only the first match claims the title; the rest are body."""
        title_shape = _StubShape("Slide Title")
        slide = _StubSlide(
            _StubShapes([title_shape, _StubShape("More")], title=title_shape)
        )

        data = _slide_to_data(slide, 1)

        assert data.title == "Slide Title"
        assert data.body == "More"

    def test_notes_are_captured(self) -> None:
        """A populated notes slide lands on ``notes``."""
        slide = _StubSlide(_StubShapes([], title=None), notes="  Speaker note  ")

        assert _slide_to_data(slide, 1).notes == "Speaker note"

    def test_malformed_notes_slide_yields_empty_notes(self) -> None:
        """A notes slide missing ``notes_text_frame`` degrades to ``""``."""
        slide = _StubSlide(_StubShapes([], title=None), notes_raise=True)

        assert _slide_to_data(slide, 1).notes == ""

    def test_no_notes_slide_yields_empty_notes(self) -> None:
        """No notes slide at all is also ``""``, not ``None``."""
        slide = _StubSlide(_StubShapes([], title=None))

        assert _slide_to_data(slide, 1).notes == ""

    def test_slide_missing_has_notes_slide_attribute_yields_empty_notes(self) -> None:
        """Undeclared notes are not read, even when they are present.

        The notes-side twin of
        ``test_shape_missing_has_text_frame_attribute_is_skipped``.
        :class:`_BareSlide` carries real notes text behind an absent
        ``has_notes_slide``, so flipping the ``getattr`` default to
        ``True`` surfaces that text and fails here rather than being
        absorbed by the ``AttributeError`` arm.
        """
        slide = _BareSlide(_StubShapes([_StubShape("Body")], title=None))

        data = _slide_to_data(slide, 1)

        assert data.notes == ""
        assert data.body == "Body"


class TestPresentationMarkdown:
    """Every combination of title/body/notes renders without stray blanks."""

    @staticmethod
    def _render(slide: SlideData) -> str:
        """Render a one-slide deck to markdown."""
        fragment = ParsedFragment(
            content="",
            metadata={"title": "Deck"},
            source_path="/tmp/deck.pptx",  # not touched, identity only
            timestamp=datetime(2024, 3, 15, tzinfo=UTC),
            payload=PresentationData(title="Deck", slides=(slide,)),
        )
        return PresentationIngestor().convert_to_markdown(fragment)

    def test_slide_with_notes_but_no_body(self) -> None:
        """A notes-only slide still emits its notes section."""
        result = self._render(
            SlideData(index=1, title="T", body="", notes="Just a note")
        )

        assert "## Slide 1: T" in result
        assert "**Speaker notes:**" in result
        assert "Just a note" in result

    def test_slide_with_body_but_no_notes(self) -> None:
        """A body-only slide emits no speaker-notes section."""
        result = self._render(SlideData(index=1, title="T", body="Body", notes=""))

        assert "Body" in result
        assert "**Speaker notes:**" not in result

    def test_slide_without_title(self) -> None:
        """A titleless slide gets a bare ``## Slide N`` heading."""
        result = self._render(SlideData(index=2, title=None, body="Body", notes=""))

        assert "## Slide 2\n" in result
        assert "## Slide 2:" not in result

    def test_empty_slide_renders_only_its_heading(self) -> None:
        """A wholly empty slide contributes just the heading line."""
        result = self._render(SlideData(index=1, title=None, body="", notes=""))

        assert result.strip().endswith("## Slide 1")


# ---- Sheet decomposition ------------------------------------------------


class TestCellCoercion:
    """Cell values are coerced to display strings without ``"None"`` leaking."""

    @pytest.mark.parametrize(
        ("value", "expected"),
        [
            pytest.param(None, "", id="none-becomes-empty-not-the-string-None"),
            pytest.param(
                datetime(2024, 3, 15, 8, 30, tzinfo=UTC),
                "2024-03-15T08:30:00+00:00",
                id="datetime-becomes-isoformat",
            ),
            pytest.param(42, "42", id="int"),
            pytest.param(3.5, "3.5", id="float"),
            pytest.param(True, "True", id="bool"),
            pytest.param("text", "text", id="str-passthrough"),
        ],
    )
    def test_cell_coercion(self, value: object, expected: str) -> None:
        """Each supported cell type renders to its display form."""
        assert _cell_to_str(value) == expected


class TestHeaderSplitting:
    """Header auto-detection and its two explicit overrides (#165)."""

    def test_empty_sheet_has_no_headers(self) -> None:
        """No rows means no first row to promote."""
        sheet = _split_header("Sheet1", [])
        assert sheet.headers is None
        assert sheet.rows == ()

    def test_empty_sheet_with_forced_header_still_has_none(self) -> None:
        """``has_header=True`` cannot invent a header out of nothing."""
        sheet = _split_header("Sheet1", [], has_header=True)
        assert sheet.headers is None

    def test_blank_cell_in_first_row_defeats_auto_detection(self) -> None:
        """A first row with a blank cell is data, not a header."""
        sheet = _split_header("Sheet1", [("a", "  "), ("1", "2")])
        assert sheet.headers is None
        assert sheet.rows == (("a", "  "), ("1", "2"))

    def test_all_non_blank_first_row_is_auto_promoted(self) -> None:
        """The positive arm of the heuristic."""
        sheet = _split_header("Sheet1", [("a", "b"), ("1", "2")])
        assert sheet.headers == ("a", "b")
        assert sheet.rows == (("1", "2"),)

    def test_forced_header_overrides_the_heuristic(self) -> None:
        """``has_header=True`` promotes even a blank-bearing first row."""
        sheet = _split_header("Sheet1", [("a", ""), ("1", "2")], has_header=True)
        assert sheet.headers == ("a", "")

    def test_forced_no_header_overrides_the_heuristic(self) -> None:
        """``has_header=False`` demotes even a clean first row."""
        sheet = _split_header("Sheet1", [("a", "b"), ("1", "2")], has_header=False)
        assert sheet.headers is None
        assert sheet.rows == (("a", "b"), ("1", "2"))


class TestSpreadsheetMarkdown:
    """Rendering copes with an empty sheet and with a missing payload."""

    @staticmethod
    def _fragment(payload: object) -> ParsedFragment:
        """Build a one-sheet spreadsheet fragment around *payload*."""
        return ParsedFragment(
            content="",
            metadata={"original_file": "book.xlsx"},
            source_path="/tmp/book.xlsx",  # not touched, identity only
            timestamp=datetime(2024, 3, 15, tzinfo=UTC),
            payload=payload,
        )

    def test_empty_sheet_renders_a_placeholder(self) -> None:
        """An empty sheet says so rather than emitting a headerless table."""
        fragment = self._fragment(SheetData(name="Sheet1", headers=None, rows=()))

        result = SpreadsheetIngestor().convert_to_markdown(fragment)

        assert "_(empty sheet)_" in result
        assert "|" not in result

    def test_missing_payload_renders_as_empty(self) -> None:
        """A fragment whose payload is not a ``SheetData`` yields no rows."""
        fragment = self._fragment(None)

        result = SpreadsheetIngestor().convert_to_markdown(fragment)

        assert "_(empty sheet)_" in result

    def test_headerless_rows_get_generated_column_names(self) -> None:
        """Rows without headers render under ``col1``/``col2`` placeholders."""
        fragment = self._fragment(
            SheetData(name="Sheet1", headers=None, rows=(("1", "2"),))
        )

        result = SpreadsheetIngestor().convert_to_markdown(fragment)

        assert "col1" in result
        assert "col2" in result
