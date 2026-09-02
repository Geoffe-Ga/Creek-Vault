"""Tests for creek.ingest.documents — Document file ingestor.

Covers discovery of DOCX/PDF/HTML/TXT files, format-specific parsing,
markdown conversion, frontmatter generation, scanned PDF detection,
metadata extraction, and edge cases.
"""

from __future__ import annotations

from datetime import datetime
from typing import TYPE_CHECKING, Any
from unittest.mock import patch
from zoneinfo import ZoneInfo

if TYPE_CHECKING:
    from pathlib import Path

import pytest

from creek.ingest.base import (
    ParsedFragment,
    RawDocument,
    assemble_ingested_fragment,
)
from creek.ingest.documents import (
    DocumentIngestor,
    _detect_scanned_pdf,
    _extract_docx_metadata,
    _extract_pdf_metadata,
    _infer_document_platform,
    _parse_docx_to_markdown,
    _parse_html_to_markdown,
    _parse_pdf_to_text,
    _wrap_txt_with_structure,
)
from creek.models import Authorship, SourcePlatform

LA_TZ = ZoneInfo("America/Los_Angeles")

# ---- Supported extensions ----

DOCUMENT_EXTENSIONS = (".docx", ".pdf", ".html", ".htm", ".txt", ".rtf")


# ---- Fixtures ----


@pytest.fixture()
def doc_ingestor() -> DocumentIngestor:
    """Create a DocumentIngestor instance for testing."""
    return DocumentIngestor()


@pytest.fixture()
def tmp_doc_dir(tmp_path: Path) -> Path:
    """Create a temporary directory with sample document files.

    Returns:
        Path to the temporary directory containing test files.
    """
    # HTML file
    html_file = tmp_path / "sample.html"
    html_file.write_text(
        "<html><body><h1>Hello World</h1>"
        "<p>This is a <strong>test</strong> document.</p>"
        "<ul><li>Item 1</li><li>Item 2</li></ul>"
        "</body></html>",
        encoding="utf-8",
    )

    # HTM file
    htm_file = tmp_path / "sample.htm"
    htm_file.write_text(
        "<html><body><h1>HTM File</h1><p>Content here.</p></body></html>",
        encoding="utf-8",
    )

    # TXT file
    txt_file = tmp_path / "sample.txt"
    txt_file.write_text(
        "Meeting Notes\n\nDiscussed the project timeline.\n"
        "Action items:\n- Complete design review\n- Update documentation\n",
        encoding="utf-8",
    )

    # Non-document file (should be ignored)
    py_file = tmp_path / "script.py"
    py_file.write_text("print('hello')", encoding="utf-8")

    # Nested directory with HTML
    subdir = tmp_path / "sub"
    subdir.mkdir()
    nested_html = subdir / "nested.html"
    nested_html.write_text(
        "<html><body><p>Nested content.</p></body></html>",
        encoding="utf-8",
    )

    return tmp_path


def _make_raw_doc(
    path: Path,
    content: bytes,
    metadata: dict[str, Any] | None = None,
    encoding: str = "utf-8",
) -> RawDocument:
    """Create a RawDocument for testing.

    Args:
        path: Filesystem path.
        content: Raw byte content.
        metadata: Optional metadata dict.
        encoding: Detected encoding.

    Returns:
        A RawDocument instance.
    """
    return RawDocument(
        path=path,
        content=content,
        metadata=metadata or {"source_type": "document"},
        detected_encoding=encoding,
    )


def _make_fragment(
    content: str,
    source_path: str,
    metadata: dict[str, Any] | None = None,
) -> ParsedFragment:
    """Create a ParsedFragment for testing.

    Args:
        content: Text content.
        source_path: Path string to source.
        metadata: Optional metadata dict.

    Returns:
        A ParsedFragment instance.
    """
    return ParsedFragment(
        content=content,
        metadata=metadata or {},
        source_path=source_path,
        timestamp=datetime(2024, 1, 15, 10, 30, tzinfo=LA_TZ),
    )


# ---- Discovery Tests ----


class TestDocumentIngestorDiscover:
    """Tests for DocumentIngestor.discover()."""

    def test_discover_empty_directory(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Discover returns empty list for directory with no documents."""
        result = doc_ingestor.discover(tmp_path)
        assert result == []

    def test_discover_nonexistent_path(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Discover returns empty list for nonexistent path."""
        result = doc_ingestor.discover(tmp_path / "nonexistent")
        assert result == []

    def test_discover_finds_html_files(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover finds .html and .htm files."""
        results = doc_ingestor.discover(tmp_doc_dir)
        extensions = {r.path.suffix for r in results}
        assert ".html" in extensions
        assert ".htm" in extensions

    def test_discover_finds_txt_files(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover finds .txt files."""
        results = doc_ingestor.discover(tmp_doc_dir)
        extensions = {r.path.suffix for r in results}
        assert ".txt" in extensions

    def test_discover_ignores_non_document_files(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover skips files with unsupported extensions."""
        results = doc_ingestor.discover(tmp_doc_dir)
        extensions = {r.path.suffix for r in results}
        assert ".py" not in extensions

    def test_discover_finds_nested_files(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover recurses into subdirectories."""
        results = doc_ingestor.discover(tmp_doc_dir)
        nested = [r for r in results if "sub" in str(r.path)]
        assert len(nested) >= 1

    def test_discover_single_file(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover handles a single file path."""
        html_file = tmp_doc_dir / "sample.html"
        results = doc_ingestor.discover(html_file)
        assert len(results) == 1
        assert results[0].path == html_file

    def test_discover_sets_metadata(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover sets source_type and file_type metadata."""
        results = doc_ingestor.discover(tmp_doc_dir)
        for raw_doc in results:
            assert raw_doc.metadata["source_type"] == "document"
            assert "file_type" in raw_doc.metadata

    def test_discover_detects_encoding(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Discover detects character encoding for each file."""
        results = doc_ingestor.discover(tmp_doc_dir)
        for raw_doc in results:
            assert raw_doc.detected_encoding != ""


# ---- HTML Parsing Tests ----


class TestParseHtmlToMarkdown:
    """Tests for HTML to markdown conversion helper."""

    def test_html_headings_preserved(self) -> None:
        """HTML headings convert to ATX-style markdown headings."""
        html = "<h1>Title</h1><h2>Subtitle</h2>"
        result = _parse_html_to_markdown(html)
        assert "# Title" in result
        assert "## Subtitle" in result

    def test_html_emphasis_preserved(self) -> None:
        """HTML bold and italic convert to markdown emphasis."""
        html = "<p>This is <strong>bold</strong> and <em>italic</em>.</p>"
        result = _parse_html_to_markdown(html)
        assert "**bold**" in result
        assert "*italic*" in result

    def test_html_lists_preserved(self) -> None:
        """HTML lists convert to markdown lists."""
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        result = _parse_html_to_markdown(html)
        assert "Item 1" in result
        assert "Item 2" in result

    def test_html_tables_preserved(self) -> None:
        """HTML tables convert to markdown tables."""
        html = (
            "<table><tr><th>Name</th><th>Value</th></tr>"
            "<tr><td>A</td><td>1</td></tr></table>"
        )
        result = _parse_html_to_markdown(html)
        assert "Name" in result
        assert "Value" in result

    def test_html_empty_input(self) -> None:
        """Empty HTML returns empty string."""
        result = _parse_html_to_markdown("")
        assert result.strip() == ""

    def test_html_links_preserved(self) -> None:
        """HTML links convert to markdown links."""
        html = '<p><a href="https://example.com">Click here</a></p>'
        result = _parse_html_to_markdown(html)
        assert "Click here" in result

    def test_missing_markdownify_raises_helpful_error(self) -> None:
        """When the optional ``markdownify`` extra is absent, raise a clear ImportError.

        Users who installed ``creek-tools`` without the ``[documents]``
        extra should see a one-liner pointing at the install command,
        not an opaque ``ModuleNotFoundError``.
        """
        import sys

        from creek.ingest.html import parse_html_to_markdown

        with (
            patch.dict(sys.modules, {"markdownify": None}),
            pytest.raises(ImportError, match=r"creek-tools\[documents\]"),
        ):
            parse_html_to_markdown("<p>x</p>")


# ---- TXT Parsing Tests ----


class TestWrapTxtWithStructure:
    """Tests for TXT structure detection helper."""

    def test_txt_basic_wrapping(self) -> None:
        """Plain text is returned as-is content."""
        text = "Hello, this is plain text."
        result = _wrap_txt_with_structure(text)
        assert "Hello, this is plain text." in result

    def test_txt_title_detection(self) -> None:
        """First non-empty line treated as potential title."""
        text = "Meeting Notes\n\nSome content here."
        result = _wrap_txt_with_structure(text)
        assert "Meeting Notes" in result

    def test_txt_list_detection(self) -> None:
        """Lines starting with - are preserved as list items."""
        text = "Items:\n- First\n- Second\n- Third"
        result = _wrap_txt_with_structure(text)
        assert "- First" in result
        assert "- Second" in result

    def test_txt_empty_input(self) -> None:
        """Empty text returns empty string."""
        result = _wrap_txt_with_structure("")
        assert result == ""

    def test_txt_whitespace_only(self) -> None:
        """Whitespace-only text returns empty string."""
        result = _wrap_txt_with_structure("   \n  \n  ")
        assert result == ""


# ---- PDF Tests ----


class TestParsePdfToText:
    """Tests for PDF text extraction helper."""

    def test_pdf_extract_text_from_bytes(self, tmp_path: Path) -> None:
        """PDF text extraction returns string content."""
        # Mock pdfminer since we can't easily create real PDFs in tests
        with patch(
            "creek.ingest.documents._extract_pdf_text_from_bytes",
            return_value="Extracted PDF content.",
        ):
            result = _parse_pdf_to_text(b"fake-pdf-bytes")
            assert result == "Extracted PDF content."


class TestDetectScannedPdf:
    """Tests for scanned PDF detection."""

    def test_short_text_multi_page_is_scanned(self) -> None:
        """PDF with <100 chars text and multiple pages flagged as scanned."""
        assert _detect_scanned_pdf("short", page_count=5) is True

    def test_long_text_is_not_scanned(self) -> None:
        """PDF with substantial text is not flagged as scanned."""
        long_text = "a" * 200
        assert _detect_scanned_pdf(long_text, page_count=5) is False

    def test_single_page_short_text_not_scanned(self) -> None:
        """Single-page PDF with short text is not flagged as scanned."""
        assert _detect_scanned_pdf("short", page_count=1) is False

    def test_empty_text_multi_page_is_scanned(self) -> None:
        """PDF with no text and multiple pages flagged as scanned."""
        assert _detect_scanned_pdf("", page_count=3) is True


# ---- DOCX Tests ----


class TestParseDocxToMarkdown:
    """Tests for DOCX to markdown conversion helper."""

    def test_docx_paragraphs_converted(self, tmp_path: Path) -> None:
        """DOCX paragraphs convert to markdown text."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Second paragraph.")

        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))
        raw_bytes = docx_path.read_bytes()

        result = _parse_docx_to_markdown(raw_bytes)
        assert "Test Document" in result
        assert "test paragraph" in result
        assert "Second paragraph" in result

    def test_docx_tables_converted(self, tmp_path: Path) -> None:
        """DOCX tables convert to markdown tables."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header1"
        table.rows[0].cells[1].text = "Header2"
        table.rows[1].cells[0].text = "Cell1"
        table.rows[1].cells[1].text = "Cell2"

        docx_path = tmp_path / "table.docx"
        doc.save(str(docx_path))
        raw_bytes = docx_path.read_bytes()

        result = _parse_docx_to_markdown(raw_bytes)
        assert "Header1" in result
        assert "Cell1" in result

    def test_docx_headings_preserved(self, tmp_path: Path) -> None:
        """DOCX headings convert to markdown heading levels."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("H1 Title", level=1)
        doc.add_heading("H2 Subtitle", level=2)
        doc.add_paragraph("Body text.")

        docx_path = tmp_path / "headings.docx"
        doc.save(str(docx_path))
        raw_bytes = docx_path.read_bytes()

        result = _parse_docx_to_markdown(raw_bytes)
        assert "# H1 Title" in result
        assert "## H2 Subtitle" in result


# ---- DOCX Metadata Tests ----


class TestExtractDocxMetadata:
    """Tests for DOCX metadata extraction."""

    def test_extract_author_and_title(self, tmp_path: Path) -> None:
        """Extract author and title from DOCX core properties."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Title"
        doc.add_paragraph("Content.")

        docx_path = tmp_path / "meta.docx"
        doc.save(str(docx_path))
        raw_bytes = docx_path.read_bytes()

        metadata = _extract_docx_metadata(raw_bytes)
        assert metadata["author"] == "Test Author"
        assert metadata["title"] == "Test Title"

    def test_missing_metadata_returns_empty_fields(self, tmp_path: Path) -> None:
        """DOCX without metadata returns dict with empty/None fields."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Content.")

        docx_path = tmp_path / "no_meta.docx"
        doc.save(str(docx_path))
        raw_bytes = docx_path.read_bytes()

        metadata = _extract_docx_metadata(raw_bytes)
        assert isinstance(metadata, dict)


# ---- PDF Metadata Tests ----


class TestExtractPdfMetadata:
    """Tests for PDF metadata extraction."""

    def test_returns_dict(self) -> None:
        """PDF metadata extraction returns a dict."""
        with patch(
            "creek.ingest.documents._extract_pdf_metadata_from_bytes",
            return_value={"author": "PDF Author", "title": "PDF Title"},
        ):
            result = _extract_pdf_metadata(b"fake-pdf")
            assert isinstance(result, dict)


# ---- Platform Inference Tests ----


class TestInferDocumentPlatform:
    """Tests for document platform inference."""

    def test_docx_platform(self) -> None:
        """DOCX files get DOCUMENT platform."""
        result = _infer_document_platform(".docx")
        assert result == SourcePlatform.DOCUMENT

    def test_pdf_platform(self) -> None:
        """PDF files get DOCUMENT platform."""
        result = _infer_document_platform(".pdf")
        assert result == SourcePlatform.DOCUMENT

    def test_html_platform(self) -> None:
        """HTML files get OTHER platform."""
        result = _infer_document_platform(".html")
        assert result == SourcePlatform.OTHER

    def test_txt_platform(self) -> None:
        """TXT files get OTHER platform."""
        result = _infer_document_platform(".txt")
        assert result == SourcePlatform.OTHER


# ---- Parse Method Tests ----


class TestDocumentIngestorParse:
    """Tests for DocumentIngestor.parse()."""

    def test_parse_html_document(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Parse returns fragment for HTML document."""
        html_file = tmp_path / "test.html"
        html_file.write_text(
            "<html><body><h1>Title</h1><p>Content.</p></body></html>",
            encoding="utf-8",
        )
        raw = _make_raw_doc(
            html_file,
            html_file.read_bytes(),
            {"source_type": "document", "file_type": ".html"},
        )
        fragments = doc_ingestor.parse(raw)
        assert len(fragments) == 1
        assert "Content" in fragments[0].content

    def test_parse_txt_document(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Parse returns fragment for TXT document."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello world.\nSecond line.", encoding="utf-8")
        raw = _make_raw_doc(
            txt_file,
            txt_file.read_bytes(),
            {"source_type": "document", "file_type": ".txt"},
        )
        fragments = doc_ingestor.parse(raw)
        assert len(fragments) == 1
        assert "Hello world" in fragments[0].content

    def test_parse_docx_document(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Parse returns fragment for DOCX document."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("DOCX Title", level=1)
        doc.add_paragraph("DOCX content here.")

        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))
        raw = _make_raw_doc(
            docx_path,
            docx_path.read_bytes(),
            {"source_type": "document", "file_type": ".docx"},
        )
        fragments = doc_ingestor.parse(raw)
        assert len(fragments) == 1
        assert "DOCX content" in fragments[0].content

    def test_parse_sets_metadata(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Parse sets file_type and source_encoding in fragment metadata."""
        txt_file = tmp_path / "meta.txt"
        txt_file.write_text("Content.", encoding="utf-8")
        raw = _make_raw_doc(
            txt_file,
            txt_file.read_bytes(),
            {"source_type": "document", "file_type": ".txt"},
        )
        fragments = doc_ingestor.parse(raw)
        assert fragments[0].metadata["file_type"] == ".txt"
        assert "source_encoding" in fragments[0].metadata


# ---- Convert to Markdown Tests ----


class TestDocumentIngestorConvert:
    """Tests for DocumentIngestor.convert_to_markdown()."""

    def test_convert_returns_content(self, doc_ingestor: DocumentIngestor) -> None:
        """Convert returns the fragment content as markdown."""
        fragment = _make_fragment("# Title\n\nSome content.", "/path/test.html")
        result = doc_ingestor.convert_to_markdown(fragment)
        assert "# Title" in result
        assert "Some content." in result

    def test_convert_returns_string(self, doc_ingestor: DocumentIngestor) -> None:
        """Convert always returns a string."""
        fragment = _make_fragment("text", "/path/test.txt")
        result = doc_ingestor.convert_to_markdown(fragment)
        assert isinstance(result, str)


# ---- Frontmatter Generation Tests ----


class TestDocumentIngestorFrontmatter:
    """Tests for DocumentIngestor.generate_frontmatter()."""

    def test_frontmatter_has_required_fields(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """Frontmatter contains type, source, and created fields."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {"file_type": ".docx", "source_encoding": "utf-8"},
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["type"] == "fragment"
        assert "source" in fm
        assert "created" in fm

    def test_frontmatter_source_platform(self, doc_ingestor: DocumentIngestor) -> None:
        """Frontmatter source.platform set based on file type."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {"file_type": ".docx", "source_encoding": "utf-8"},
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["source"]["platform"] == SourcePlatform.DOCUMENT

    def test_frontmatter_original_file(self, doc_ingestor: DocumentIngestor) -> None:
        """Frontmatter preserves source.original_file."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.pdf",
            {"file_type": ".pdf", "source_encoding": "utf-8"},
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["source"]["original_file"] == "/path/test.pdf"

    def test_frontmatter_original_encoding(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """Frontmatter records source.original_encoding."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.txt",
            {"file_type": ".txt", "source_encoding": "latin-1"},
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["source"]["original_encoding"] == "latin-1"

    def test_frontmatter_includes_author_from_metadata(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """A free-text author name is preserved under ``source.author_name``."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {
                "file_type": ".docx",
                "source_encoding": "utf-8",
                "author": "Jane Doe",
            },
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["source"]["author_name"] == "Jane Doe"

    def test_frontmatter_includes_title_from_metadata(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """Frontmatter uses metadata title when available."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {
                "file_type": ".docx",
                "source_encoding": "utf-8",
                "title": "My Document",
            },
        )
        fm = doc_ingestor.generate_frontmatter(fragment)
        assert fm["title"] == "My Document"

    def test_the_scanned_flag_survives_assembly_instead_of_being_dropped(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """The scanned marker must reach the writer, not just the generator.

        This test used to call ``generate_frontmatter`` directly and assert
        ``fm["source"]["scanned"] is True``. It passed — and the key reached
        **no vault file whatsoever**, because ``FragmentSource`` does not
        model ``scanned``, ``Fragment`` leaves pydantic's ``extra="ignore"``
        in place, and ``PASSTHROUGH_FRONTMATTER_KEYS`` is a *top-level*
        allowlist that cannot rescue a key nested under ``source``. Measured
        at HEAD before #1639: an operator's scanned PDF produced an empty
        fragment carrying no indication of any kind.

        So the assertion now runs the real assembly chokepoint,
        ``assemble_ingested_fragment``, which is where the drop happened, and
        reads ``extra_frontmatter`` — the seam that actually carries a key to
        the writer. The end-to-end proof on written bytes lives in
        ``tests/test_ocr_config_wiring.py``.
        """
        fragment = _make_fragment(
            "Content.",
            "/path/test.pdf",
            {
                "file_type": ".pdf",
                "source_encoding": "utf-8",
                "scanned": True,
            },
        )
        fragment.metadata["markdown"] = doc_ingestor.convert_to_markdown(fragment)
        fragment.metadata["frontmatter"] = doc_ingestor.generate_frontmatter(fragment)

        assembled = assemble_ingested_fragment(fragment)

        assert assembled.extra_frontmatter["scanned"] is True

    def test_an_unscanned_document_gains_no_scanned_key_at_all(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """The marker rides on presence, so a normal PDF stays unmarked.

        The control for the test above: without it, "scanned reaches the
        writer" would be satisfied by a change that stamped every document in
        the vault as a scan.
        """
        fragment = _make_fragment(
            "Content.",
            "/path/test.pdf",
            {"file_type": ".pdf", "source_encoding": "utf-8"},
        )
        fragment.metadata["markdown"] = doc_ingestor.convert_to_markdown(fragment)
        fragment.metadata["frontmatter"] = doc_ingestor.generate_frontmatter(fragment)

        assembled = assemble_ingested_fragment(fragment)

        assert "scanned" not in assembled.extra_frontmatter


# ---- Full Pipeline Integration Tests ----


class TestDocumentIngestorIngest:
    """Integration tests for the full ingest pipeline."""

    def test_ingest_html_files(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Full ingest pipeline processes HTML files end-to-end."""
        result = doc_ingestor.ingest(tmp_doc_dir)
        assert len(result.fragments) > 0
        assert len(result.errors) == 0

    def test_ingest_single_txt_file(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Full ingest pipeline processes a single TXT file."""
        txt_file = tmp_doc_dir / "sample.txt"
        result = doc_ingestor.ingest(txt_file)
        assert len(result.fragments) == 1

    def test_ingest_empty_directory(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """Full ingest of empty directory produces no fragments."""
        result = doc_ingestor.ingest(tmp_path)
        assert len(result.fragments) == 0
        assert len(result.errors) == 0

    def test_ingest_provenance_recorded(
        self, doc_ingestor: DocumentIngestor, tmp_doc_dir: Path
    ) -> None:
        """Ingest records provenance for each processed fragment."""
        result = doc_ingestor.ingest(tmp_doc_dir)
        assert len(result.provenance) == len(result.fragments)
        for entry in result.provenance:
            assert entry.ingestor_name == "DocumentIngestor"
            assert entry.status == "success"


class TestDocumentIngestorAuthoredAtPerFormat:
    """FEAT-031: each document format extracts ``authored_at`` from its native metadata.

    HTML reads meta tags / JSON-LD, PDF reads ``/CreationDate``, DOCX
    reads ``dcterms:created``, RTF reads ``\\creatim``. TXT has no
    embedded date metadata so it must yield ``None``.
    """

    def _ingest_one(
        self, doc_ingestor: DocumentIngestor, file_path: Path
    ) -> ParsedFragment:
        """Parse one file and return its single ParsedFragment."""
        docs = doc_ingestor.discover(file_path)
        assert len(docs) == 1
        parsed = doc_ingestor.parse(docs[0])
        assert len(parsed) == 1
        return parsed[0]

    def test_html_meta_article_published_time(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """``<meta property="article:published_time">`` is the top of the chain."""
        from datetime import UTC

        html_file = tmp_path / "post.html"
        html_file.write_text(
            "<html><head>"
            '<meta property="article:published_time" '
            'content="2024-03-15T08:30:00Z">'
            '<meta name="date" content="2025-01-01">'
            "</head><body><p>Hi</p></body></html>",
            encoding="utf-8",
        )
        parsed = self._ingest_one(doc_ingestor, html_file)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        assert authored == datetime(2024, 3, 15, 8, 30, tzinfo=UTC)

    def test_html_dc_date_fallback(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """No article:published_time / OG → Dublin Core wins next."""
        from datetime import UTC

        html_file = tmp_path / "post.html"
        html_file.write_text(
            "<html><head>"
            '<meta name="DC.date.issued" content="2024-04-01">'
            "</head><body><p>Hi</p></body></html>",
            encoding="utf-8",
        )
        parsed = self._ingest_one(doc_ingestor, html_file)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        assert authored == datetime(2024, 4, 1, tzinfo=UTC)

    def test_html_json_ld_date_published(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """JSON-LD ``datePublished`` is the final HTML fallback."""
        from datetime import UTC

        html_file = tmp_path / "post.html"
        html_file.write_text(
            "<html><head>"
            '<script type="application/ld+json">'
            '{"@type": "Article", "datePublished": "2024-05-15"}'
            "</script></head><body><p>Hi</p></body></html>",
            encoding="utf-8",
        )
        parsed = self._ingest_one(doc_ingestor, html_file)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        assert authored == datetime(2024, 5, 15, tzinfo=UTC)

    def test_html_no_metadata_yields_none(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """HTML with no date metadata → ``authored_at`` is ``None``."""
        html_file = tmp_path / "post.html"
        html_file.write_text(
            "<html><body><h1>Just a page</h1></body></html>",
            encoding="utf-8",
        )
        parsed = self._ingest_one(doc_ingestor, html_file)
        assert parsed.metadata["authored_at"] is None

    def test_txt_files_always_yield_none(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """TXT files carry no embedded date metadata → ``None``."""
        txt = tmp_path / "notes.txt"
        txt.write_text("Hello\nWorld\n", encoding="utf-8")
        parsed = self._ingest_one(doc_ingestor, txt)
        assert parsed.metadata["authored_at"] is None

    def test_docx_authored_at_from_dcterms_created(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """A real DOCX's core ``dcterms:created`` lands on ``authored_at``."""
        from datetime import UTC

        from docx import Document as DocxDocument

        docx_path = tmp_path / "report.docx"
        doc = DocxDocument()
        doc.add_paragraph("Hello.")
        doc.core_properties.created = datetime(2024, 3, 15, 8, 30, tzinfo=UTC)
        doc.save(docx_path)
        parsed = self._ingest_one(doc_ingestor, docx_path)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        # python-docx strips tz on round-trip; we just need the date.
        assert authored.date() == datetime(2024, 3, 15).date()

    def test_rtf_authored_at_from_creatim(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """RTF ``\\creatim`` control word is parsed into ``authored_at``."""
        from datetime import UTC

        rtf_path = tmp_path / "doc.rtf"
        rtf_path.write_text(
            "{\\rtf1\\ansi"
            "{\\info"
            "{\\creatim\\yr2024\\mo3\\dy15\\hr8\\min30\\sec0}"
            "{\\revtim\\yr2024\\mo4\\dy1\\hr9\\min0\\sec0}"
            "}"
            "Body text.}",
            encoding="ascii",
        )
        parsed = self._ingest_one(doc_ingestor, rtf_path)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        assert authored == datetime(2024, 3, 15, 8, 30, tzinfo=UTC)

    def test_rtf_falls_back_to_revtim(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """No ``\\creatim`` → ``\\revtim`` is the next candidate."""
        from datetime import UTC

        rtf_path = tmp_path / "doc.rtf"
        rtf_path.write_text(
            "{\\rtf1\\ansi"
            "{\\info"
            "{\\revtim\\yr2024\\mo4\\dy1\\hr9\\min0\\sec0}"
            "}"
            "Body text.}",
            encoding="ascii",
        )
        parsed = self._ingest_one(doc_ingestor, rtf_path)
        authored = parsed.metadata["authored_at"]
        assert authored is not None
        assert authored == datetime(2024, 4, 1, 9, 0, tzinfo=UTC)

    def test_rtf_no_dates_yields_none(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """An RTF with no date control words → ``None``."""
        rtf_path = tmp_path / "doc.rtf"
        rtf_path.write_text(
            "{\\rtf1\\ansi Body text.}",
            encoding="ascii",
        )
        parsed = self._ingest_one(doc_ingestor, rtf_path)
        assert parsed.metadata["authored_at"] is None

    def test_authored_at_lands_in_generated_frontmatter(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """When extracted, ``authored_at`` is an ISO string in frontmatter."""
        html_file = tmp_path / "post.html"
        html_file.write_text(
            "<html><head>"
            '<meta property="article:published_time" '
            'content="2024-03-15T08:30:00Z">'
            "</head><body><p>Hi</p></body></html>",
            encoding="utf-8",
        )
        parsed = self._ingest_one(doc_ingestor, html_file)
        fm = doc_ingestor.generate_frontmatter(parsed)
        assert "authored_at" in fm
        assert fm["authored_at"].startswith("2024-03-15")

    def test_no_authored_at_omits_key_from_frontmatter(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """``None`` ``authored_at`` is absent from the frontmatter dict."""
        txt = tmp_path / "notes.txt"
        txt.write_text("Hello\n", encoding="utf-8")
        parsed = self._ingest_one(doc_ingestor, txt)
        fm = doc_ingestor.generate_frontmatter(parsed)
        assert "authored_at" not in fm


# ---- Document author metadata (issue #1229) ----


def _real_docx_bytes(path: Path, author: str) -> Path:
    """Write a genuine ``.docx`` at *path* stamped with *author*, and return it.

    Word stamps ``core_properties.author`` on every save, so a fixture without
    one is not a realistic document.
    """
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.core_properties.author = author
    document.add_paragraph("A real Word document.")
    document.save(path)
    return path


def _real_pdf_bytes(path: Path, author: str) -> Path:
    """Write a minimal but genuine PDF at *path* with ``/Author`` set.

    Hand-assembled (rather than mocked) so the ``/Info``-dictionary path in
    :func:`creek.ingest.documents._extract_pdf_metadata` is exercised for real:
    a mocked extractor could not have caught issue #1229.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>",
        b"<< /Author (" + author.encode("ascii") + b") /Title (Notes) >>",
    ]
    body = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += b"%d 0 obj\n" % number + obj + b"\nendobj\n"
    xref_offset = len(body)
    body += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        body += b"%010d 00000 n \n" % offset
    body += (
        b"trailer\n<< /Size %d /Root 1 0 R /Info 4 0 R >>\nstartxref\n%d\n%%%%EOF\n"
        % (len(objects) + 1, xref_offset)
    )
    path.write_bytes(bytes(body))
    return path


class TestDocumentAuthorMetadata:
    """A document's free-text author must never land in the Authorship enum.

    ``FragmentSource.author`` answers "whose views does this stand for?" on the
    ``self|ai|other|collaborative`` axis. A DOCX ``core_properties.author`` or a
    PDF ``/Author`` answers the different question "what name is on the file".
    Issue #1229: writing the second into the first made every real Word document
    fail Pydantic validation and get dropped at assembly.
    """

    def _assemble_one(
        self, doc_ingestor: DocumentIngestor, file_path: Path
    ) -> ParsedFragment:
        """Ingest *file_path* and return its single ``ParsedFragment``."""
        result = doc_ingestor.ingest(file_path)
        assert result.errors == []
        assert len(result.fragments) == 1
        return result.fragments[0]

    def test_word_document_with_author_ingests(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """A ``.docx`` carrying an author assembles into a valid Fragment."""
        path = _real_docx_bytes(tmp_path / "real.docx", "Jane Q. Author")
        parsed = self._assemble_one(doc_ingestor, path)

        assembled = assemble_ingested_fragment(parsed)

        assert assembled.fragment.source.author == Authorship.OTHER
        assert assembled.fragment.source.author_name == "Jane Q. Author"

    def test_pdf_with_author_ingests(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """A ``.pdf`` carrying ``/Author`` assembles into a valid Fragment."""
        path = _real_pdf_bytes(tmp_path / "real.pdf", "Ada Lovelace")
        parsed = self._assemble_one(doc_ingestor, path)

        assembled = assemble_ingested_fragment(parsed)

        assert assembled.fragment.source.author == Authorship.OTHER
        assert assembled.fragment.source.author_name == "Ada Lovelace"

    def test_rtf_and_html_carry_no_extracted_author(
        self, doc_ingestor: DocumentIngestor, tmp_path: Path
    ) -> None:
        """RTF and HTML never populate ``author``, so they keep the self default.

        Issue #1229 asserted these two shared the DOCX/PDF defect. They do not:
        ``DocumentIngestor._extract_metadata`` branches only on ``.docx`` and
        ``.pdf``, so an RTF ``\\author`` group and an HTML ``<meta name=author>``
        are never read. This pins that, so a future author-extraction branch for
        either format cannot reintroduce the enum collision unnoticed.
        """
        rtf = tmp_path / "note.rtf"
        rtf.write_text(
            "{\\rtf1\\ansi{\\info{\\author Ada Lovelace}}Body text.}",
            encoding="ascii",
        )
        html = tmp_path / "note.html"
        html.write_text(
            '<html><head><meta name="author" content="Ada Lovelace">'
            "</head><body><p>Body text.</p></body></html>",
            encoding="utf-8",
        )

        for path in (rtf, html):
            parsed = self._assemble_one(doc_ingestor, path)
            assembled = assemble_ingested_fragment(parsed)
            assert assembled.fragment.source.author == Authorship.SELF
            assert assembled.fragment.source.author_name is None

    def test_named_author_maps_to_other(self, doc_ingestor: DocumentIngestor) -> None:
        """A person's name resolves the axis to ``other``, not ``self``."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {
                "file_type": ".docx",
                "source_encoding": "utf-8",
                "author": "Jane Doe",
            },
        )

        fm = doc_ingestor.generate_frontmatter(fragment)

        assert fm["source"]["author"] == Authorship.OTHER
        assert fm["source"]["author_name"] == "Jane Doe"

    def test_valid_authorship_value_is_honoured(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """A metadata author that already names an axis member is kept as one."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {
                "file_type": ".docx",
                "source_encoding": "utf-8",
                "author": "ai",
            },
        )

        fm = doc_ingestor.generate_frontmatter(fragment)

        assert fm["source"]["author"] == Authorship.AI
        assert "author_name" not in fm["source"]

    def test_blank_author_leaves_the_self_default(
        self, doc_ingestor: DocumentIngestor
    ) -> None:
        """A whitespace-only author is no evidence, so nothing is written."""
        fragment = _make_fragment(
            "Content.",
            "/path/test.docx",
            {
                "file_type": ".docx",
                "source_encoding": "utf-8",
                "author": "   ",
            },
        )

        fm = doc_ingestor.generate_frontmatter(fragment)

        assert "author" not in fm["source"]
        assert "author_name" not in fm["source"]
