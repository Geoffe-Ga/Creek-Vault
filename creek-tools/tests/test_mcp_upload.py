"""``creek.upload`` MCP tool — Adepthood document bytes → fragment (#1023).

Drives the REAL ledger-backed ingest (``run_ingest``) against a temp vault so
the guarantees that matter are exercised end-to-end rather than mocked:

- the gate ORDER (tier → size → decode → overwrite → write) is observable from
  outside, because a refusal at an earlier gate leaves the later gate's
  side effects absent from disk;
- staging is byte-exact, so a real ``.docx`` ZIP container survives it;
- re-sending identical bytes is a true no-op (no second fragment, no mtime
  bump), and changed bytes update the same fragment in place;
- the document's bytes never reach the audit log, deliberately tested with a
  payload short enough to hit ``summarise_args``' verbatim pass-through.

Binary fixtures are generated in-process (``python-docx`` / ``openpyxl``) —
there are no checked-in binaries — and every base64 payload is derived at
runtime, never pasted as an opaque literal.
"""

from __future__ import annotations

import base64
import json
import shutil
from typing import TYPE_CHECKING, Any

import frontmatter
import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from creek.ingest import route_to_ingestor
from creek.ingest.journal_staging import UPLOAD_STAGING_RELDIR
from creek.ingest.pipeline import IngestRunResult
from creek.purge import PurgeEngine
from creek_mcp.audit import MCP_AUDIT_RELPATH
from creek_mcp.read_gate import GENERIC_ABOVE_CEILING_REASON
from creek_mcp.staged_names import MAX_SUFFIX_CHARS, safe_stem
from creek_mcp.tier_ceiling import TierCeiling
from creek_mcp.tools.upload import (
    _MAX_ENCODED_CHARS,
    MAX_UPLOAD_BYTES,
    TOOL_NAME,
    _upload_staged_path,
    upload_tool,
)

if TYPE_CHECKING:
    from pathlib import Path

_TS = "2026-06-20T10:00:00+00:00"
_REFUSAL_KEYS = {"status", "tool", "tier_ceiling", "reason"}

# Two byte sentinels, deliberately disjoint from every tier word, so an
# assertion about "intimate" and an assertion about protected plaintext can
# never accidentally satisfy each other.
_SECRET = b"synthetic-tender-secret-1023"
_BENIGN = b"benign-open-replacement-1023"


def _vault(tmp_path: Path) -> Path:
    """Create the minimum vault layout the writer + ledger + audit need."""
    for sub in ("00-Creek-Meta/State", "00-Creek-Meta/audit", "01-Fragments"):
        (tmp_path / sub).mkdir(parents=True, exist_ok=True)
    return tmp_path


def _fragments(vault: Path) -> list[Path]:
    """Return all fragment files under ``01-Fragments``."""
    return sorted((vault / "01-Fragments").rglob("*.md"))


def _staged(vault: Path) -> list[Path]:
    """Return every staged upload file under the uploads staging dir."""
    root = vault / "00-Creek-Meta" / "adepthood" / "uploads"
    return sorted(path for path in root.rglob("*") if path.is_file())


def _audit(vault: Path) -> list[dict[str, Any]]:
    """Return parsed MCP audit-log entries."""
    log = vault / MCP_AUDIT_RELPATH
    assert log.exists()
    return [json.loads(line) for line in log.read_text().splitlines() if line.strip()]


def _leaking_files(vault: Path, needle: bytes) -> list[Path]:
    """Return every file anywhere under *vault* whose bytes contain *needle*.

    Byte-based on purpose: the journal suite's text-based twin globs ``*.md``
    and would both miss a staged ``.docx`` and raise on decoding it.
    """
    return [
        path
        for path in sorted(vault.rglob("*"))
        if path.is_file() and needle in path.read_bytes()
    ]


def _b64(payload: bytes) -> str:
    """Return the base64 of *payload*, derived at runtime, never hardcoded."""
    return base64.b64encode(payload).decode("ascii")


_DOCX_AUTHOR = "Dana Rivers"
"""The author name the ``.docx`` fixture carries, as every real Word save does."""


def _docx_bytes(tmp_path: Path) -> bytes:
    """Build a genuine ``.docx`` ZIP container in-process and return its bytes.

    ``core_properties.author`` is stamped rather than cleared: Word writes one on
    every save, so a document without it is not the document users upload. It
    used to be cleared here to route around issue #1229 (the ingestor copied the
    name onto the ``Authorship`` enum, so any real document failed to assemble);
    that workaround is gone with the defect.
    """
    path = tmp_path / "source.docx"
    document = DocxDocument()
    document.core_properties.author = _DOCX_AUTHOR
    document.add_heading("Quarterly notes", level=1)
    document.add_paragraph("The creek runs steady through the autumn.")
    document.save(path)
    return path.read_bytes()


def _xlsx_bytes(tmp_path: Path) -> bytes:
    """Build a genuine ``.xlsx`` workbook in-process and return its bytes."""
    path = tmp_path / "source.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet["A1"] = "month"
    sheet["B1"] = "spend"
    sheet["A2"] = "june"
    sheet["B2"] = 42
    workbook.save(path)
    return path.read_bytes()


def _failing_runner(**_kwargs: Any) -> IngestRunResult:
    """Ingest-runner stub that reports a hard write failure."""
    return IngestRunResult(
        written=0,
        errors=["boom"],
        discovered=1,
        created=0,
        updated=0,
        unchanged=0,
        tombed=0,
        skipped=0,
    )


def test_tier_gate_refuses_before_the_bytes_are_even_decoded(tmp_path: Path) -> None:
    """An ``intimate`` upload at ``ceiling=open`` never decodes a byte.

    The payload is deliberately undecodable, so a refusal that names base64
    would prove the decode ran first — the one ordering this AC pins. The
    control call at the end sends the SAME undecodable payload at an admitted
    tier: only then does the decode run, and only then does the reason flip.
    """
    vault = _vault(tmp_path)
    undecodable = "!!!not-base64!!!"

    result = upload_tool(
        vault_path=vault,
        filename="note.txt",
        content_base64=undecodable,
        external_id="u-1",
        timestamp=_TS,
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert set(result) == _REFUSAL_KEYS
    reason = str(result["reason"])
    assert "intimate" in reason
    assert "base64" not in reason
    assert not (vault / UPLOAD_STAGING_RELDIR).exists()
    assert _fragments(vault) == []
    entries = _audit(vault)
    assert len(entries) == 1
    assert "created_tier" not in entries[0]
    assert "content_base64" not in entries[0]["args_summary"]

    admitted = upload_tool(
        vault_path=vault,
        filename="note.txt",
        content_base64=undecodable,
        external_id="u-1",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert admitted["reason"] == "content_base64 is not valid base64"
    assert not (vault / UPLOAD_STAGING_RELDIR).exists()
    # A malformed payload is a malformed argument: refused, never audited.
    assert len(_audit(vault)) == 1


@pytest.mark.parametrize("oversize_when_encoded", [True, False])
def test_oversize_upload_is_refused_before_anything_is_written(
    tmp_path: Path, *, oversize_when_encoded: bool
) -> None:
    """Both size gates refuse structurally, never by exception.

    ``oversize_when_encoded`` picks the *pre*-decode guard (an encoded string
    longer than any in-cap payload could produce) or the *post*-decode one (a
    payload one byte over the cap, whose base64 is exactly at the pre-check
    boundary and so must pass through it).
    """
    vault = _vault(tmp_path)
    payload = (
        "A" * (_MAX_ENCODED_CHARS + 4)
        if oversize_when_encoded
        else _b64(b"A" * (MAX_UPLOAD_BYTES + 1))
    )
    if not oversize_when_encoded:
        assert len(payload) <= _MAX_ENCODED_CHARS

    result = upload_tool(
        vault_path=vault,
        filename="huge.txt",
        content_base64=payload,
        external_id="u-big",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert set(result) == _REFUSAL_KEYS
    assert not (vault / UPLOAD_STAGING_RELDIR).exists()
    assert _fragments(vault) == []
    summary = _audit(vault)[-1]["args_summary"]
    assert summary["encoded_len"] == len(payload)
    assert isinstance(summary["encoded_len"], int)
    assert payload not in json.dumps(summary)


def test_document_bytes_never_enter_the_audit_log(tmp_path: Path) -> None:
    """The trail carries metadata and ids only — never the document itself.

    The payload is short enough that its base64 is under 64 characters, which
    is exactly the branch ``summarise_args`` passes through VERBATIM, so an
    implementation that handed ``content_base64`` to ``append(args=...)``
    would log the whole document here.
    """
    vault = _vault(tmp_path)
    sentinel = b"creek-upload-sentinel-1023"
    encoded = _b64(sentinel)
    assert len(encoded) < 64

    result = upload_tool(
        vault_path=vault,
        filename="memo.txt",
        content_base64=encoded,
        external_id="u-audit",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )
    assert result["status"] == "ok"

    log_bytes = (vault / MCP_AUDIT_RELPATH).read_bytes()
    assert sentinel not in log_bytes
    assert encoded.encode("ascii") not in log_bytes
    entry = _audit(vault)[-1]
    assert set(entry["args_summary"]) == {
        "filename",
        "external_id",
        "tier",
        "encoded_len",
    }
    assert entry["created_path"] == "01-Fragments"
    assert entry["created_tier"] == "open"
    assert entry["affected_fragment_ids"] == [result["fragment_id"]]


def test_a_real_docx_survives_staging_byte_identical(tmp_path: Path) -> None:
    """A genuine ``.docx`` ZIP container stages byte-for-byte and ingests."""
    vault = _vault(tmp_path / "vault")
    raw = _docx_bytes(tmp_path)

    result = upload_tool(
        vault_path=vault,
        filename="Quarterly Notes.docx",
        content_base64=_b64(raw),
        external_id="u-docx",
        timestamp=_TS,
        tier="personal",
        privacy_tier_ceiling=TierCeiling.PERSONAL,
    )

    assert result["status"] == "ok"
    assert result["source_type"] == "document"
    assert result["size_bytes"] == len(raw)
    staged = _staged(vault)
    assert len(staged) == 1
    assert staged[0].name == f"{safe_stem('u-docx')}.docx"
    assert staged[0].read_bytes() == raw
    fragments = _fragments(vault)
    assert len(fragments) == 1
    # The document's author name must not relocate the fragment out of
    # 01-Fragments/ — only an explicit author_slug does that (#1229).
    assert fragments[0].parent == vault / "01-Fragments" / "Documents"
    written = frontmatter.load(fragments[0])
    assert written["privacy_tier"] == "personal"
    assert written["source"]["author"] == "other"
    assert written["source"]["author_name"] == _DOCX_AUTHOR


def test_extension_routes_to_the_matching_source_type(tmp_path: Path) -> None:
    """The staged name keeps the (lower-cased) extension, which picks the ingestor.

    Pure: no vault, no bytes — only the staged-path derivation feeding
    ``route_to_ingestor``. The cases at the end cover the caller-controlled
    edges: injectivity (two ids that slug identically must not collide), an
    id with no slug-able characters at all, and an absurd extension.
    """
    cases = [
        ("notes.md", ".md", "markdown"),
        ("memo.docx", ".docx", "document"),
        ("scan.pdf", ".pdf", "document"),
        ("readme.txt", ".txt", "document"),
        ("budget.xlsx", ".xlsx", "spreadsheet"),
        ("deck.pptx", ".pptx", "presentation"),
        ("SHOT.PNG", ".png", "image"),
        ("mystery.xyz", ".xyz", "generic"),
        ("README", "", "generic"),
    ]
    for filename, suffix, source_type in cases:
        staged = _upload_staged_path(tmp_path, "u-route", filename)
        assert staged.suffix == suffix, filename
        assert route_to_ingestor(staged) == source_type, filename

    slashed = _upload_staged_path(tmp_path, "a/b", "note.txt")
    dashed = _upload_staged_path(tmp_path, "a-b", "note.txt")
    assert slashed != dashed

    unsluggable = _upload_staged_path(tmp_path, "///", "note.txt")
    assert unsluggable.stem == safe_stem("///")
    assert unsluggable.stem.isalnum()

    absurd = _upload_staged_path(tmp_path, "u-route", "x." + "z" * 200)
    assert absurd.suffix == "." + "z" * (MAX_SUFFIX_CHARS - 1)


def test_identical_resend_is_a_true_no_op(tmp_path: Path) -> None:
    """Identical bytes under the same id neither restage nor mint a fragment.

    The ``st_mtime_ns`` assertion is the sharp one: because
    ``generate_fragment_id`` hashes the mtime-derived timestamp, an
    unconditional ``write_bytes`` would mint a second fragment here.
    """
    vault = _vault(tmp_path)
    payload = _b64(b"the creek runs steady, and keeps running steady")

    def _send() -> dict[str, Any]:
        """Send the same bytes under the same external id."""
        return upload_tool(
            vault_path=vault,
            filename="steady.txt",
            content_base64=payload,
            external_id="u-noop",
            tier="open",
            timestamp=_TS,
            privacy_tier_ceiling=TierCeiling.OPEN,
        )

    first = _send()
    assert first["action"] == "created"
    staged = _upload_staged_path(vault, "u-noop", "steady.txt")
    mtime_before = staged.stat().st_mtime_ns

    second = _send()

    assert second["action"] == "unchanged"
    assert second["fragment_id"] == first["fragment_id"]
    assert staged.stat().st_mtime_ns == mtime_before
    assert len(_fragments(vault)) == 1


def test_changed_bytes_update_the_same_fragment_in_place(tmp_path: Path) -> None:
    """A re-send with new bytes rewrites the fragment, never orphaning a twin."""
    vault = _vault(tmp_path)

    def _send(text: str) -> dict[str, Any]:
        """Send *text* as a ``.txt`` upload under a fixed external id."""
        return upload_tool(
            vault_path=vault,
            filename="draft.txt",
            content_base64=_b64(text.encode("utf-8")),
            external_id="u-edit",
            tier="open",
            timestamp=_TS,
            privacy_tier_ceiling=TierCeiling.OPEN,
        )

    first = _send("the first honest sentence about the creek")
    assert first["action"] == "created"

    second = _send("the second honest sentence about the river")

    assert second["action"] == "updated"
    assert second["fragment_id"] == first["fragment_id"]
    fragments = _fragments(vault)
    assert len(fragments) == 1
    body = fragments[0].read_text(encoding="utf-8")
    assert "second honest sentence" in body
    assert "first honest sentence" not in body


def test_overwriting_an_intimate_upload_is_refused_at_ceiling_open(
    tmp_path: Path,
) -> None:
    """You may only overwrite what you could have read (#970), and the gate
    sits ABOVE staging, so the intimate bytes are still on disk after it fires.
    """
    vault = _vault(tmp_path)
    seeded = upload_tool(
        vault_path=vault,
        filename="confession.txt",
        content_base64=_b64(b"a tender confession " + _SECRET),
        external_id="u-970",
        timestamp=_TS,
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.ALL,
    )
    assert seeded["action"] == "created"
    staged = _upload_staged_path(vault, "u-970", "confession.txt")
    assert _SECRET in staged.read_bytes()
    fragment_before = _fragments(vault)[0].read_bytes()

    result = upload_tool(
        vault_path=vault,
        filename="confession.txt",
        content_base64=_b64(b"a bland note " + _BENIGN),
        external_id="u-970",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert result["reason"] == GENERIC_ABOVE_CEILING_REASON
    assert set(result) == _REFUSAL_KEYS
    assert _SECRET in staged.read_bytes()
    assert _fragments(vault)[0].read_bytes() == fragment_before
    assert _leaking_files(vault, _BENIGN) == []


@pytest.mark.parametrize(
    ("filename", "payload"),
    [
        ("mystery.bin", b"\x00\x01binary\x00payload\x00"),
        ("export.json", b'{"claimed": "by the chat ingestors"}'),
    ],
)
def test_a_file_that_produces_no_fragment_is_refused_not_reported_ok(
    tmp_path: Path, filename: str, payload: bytes
) -> None:
    """``written == 0`` is a refusal, not a success.

    Both cases route to ``generic`` and yield ``written=0, errors=[]`` — the
    binary heuristic drops one and the claimed-``.json`` guard drops the
    other — so only inspecting ``written`` tells them apart from success.
    """
    vault = _vault(tmp_path)

    result = upload_tool(
        vault_path=vault,
        filename=filename,
        content_base64=_b64(payload),
        external_id=f"u-{filename}",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert set(result) == _REFUSAL_KEYS
    assert "no fragment" in str(result["reason"])
    assert _fragments(vault) == []


def test_ingest_failure_is_a_structured_refusal(tmp_path: Path) -> None:
    """Write failures and a missing vault both refuse, never raise."""
    vault = _vault(tmp_path)

    failed = upload_tool(
        vault_path=vault,
        filename="note.txt",
        content_base64=_b64(b"a note that never lands"),
        external_id="u-fail",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
        run=_failing_runner,
    )

    assert failed["status"] == "refused"
    assert "ingest failed: boom" in str(failed["reason"])
    assert _audit(vault)[-1]["created_tier"] == "open"

    shutil.rmtree(vault)
    gone = upload_tool(
        vault_path=vault,
        filename="note.txt",
        content_base64=_b64(b"a note with nowhere to land"),
        external_id="u-gone",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert gone["status"] == "refused"
    assert gone["reason"] == "vault unavailable"


@pytest.mark.parametrize(
    ("filename", "external_id", "content_base64", "tier"),
    [
        ("", "u-bad", "Zm9v", "open"),
        ("note.txt", "  ", "Zm9v", "open"),
        ("note.txt", "u-bad", "", "open"),
        ("note.txt", "u-bad", "Zm9v", "secret"),
    ],
)
def test_malformed_arguments_refuse_without_an_audit_entry(
    tmp_path: Path,
    filename: str,
    external_id: str,
    content_base64: str,
    tier: str,
) -> None:
    """A malformed call earns a refusal and no audit entry — there is no tier yet."""
    vault = _vault(tmp_path)

    result = upload_tool(
        vault_path=vault,
        filename=filename,
        content_base64=content_base64,
        external_id=external_id,
        tier=tier,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert result["tool"] == TOOL_NAME
    assert set(result) == _REFUSAL_KEYS
    assert _fragments(vault) == []
    assert not (vault / UPLOAD_STAGING_RELDIR).exists()
    assert not (vault / MCP_AUDIT_RELPATH).exists()


def test_omitted_tier_is_refused_before_anything_is_written(tmp_path: Path) -> None:
    """An OMITTED ``tier`` is a malformed call too, not an ``open`` upload (#1494).

    ``tier=`` is left out of the call entirely rather than passed as ``None``,
    deliberately: omission is what a real caller does, and it is what makes
    this single test kill two mutations — restoring ``tier: str = "open"`` on
    :func:`~creek_mcp.tools.upload.upload_tool`, and deleting the ``tier is
    None`` guard. Passing ``tier=None`` explicitly would kill only the second.

    The reason substring is required rather than decorative: ``PrivacyTier``
    raises ``ValueError`` on ``None``, so with the guard gone this call still
    answers ``status: refused`` — with reason ``unknown tier None`` — and an
    assertion on the status alone would be blind to the whole fix.

    The staging dir is asserted absent because an upload that got past this
    gate would have written the caller's bytes to disk before any later
    refusal could fire, which is the leak the tier default caused.
    """
    vault = _vault(tmp_path)

    result = upload_tool(
        vault_path=vault,
        filename="notes.md",
        content_base64=_b64(b"# Note\n\nA document whose tier nobody named.\n"),
        external_id="u-no-tier",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert result["tool"] == TOOL_NAME
    assert set(result) == _REFUSAL_KEYS
    assert "tier is required" in str(result["reason"])
    assert _fragments(vault) == []
    assert not (vault / UPLOAD_STAGING_RELDIR).exists()
    assert not (vault / MCP_AUDIT_RELPATH).exists()


def test_staging_path_and_origin_key_are_stable(tmp_path: Path) -> None:
    """The staging dir and the ledger key it produces are pinned (#845/#1023).

    Relocating :data:`UPLOAD_STAGING_RELDIR` without a ledger migration would
    orphan every staged upload, so its literal value is asserted here too.
    """
    assert UPLOAD_STAGING_RELDIR.as_posix() == "00-Creek-Meta/adepthood/uploads"
    vault = _vault(tmp_path / "vault")

    result = upload_tool(
        vault_path=vault,
        filename="budget.xlsx",
        content_base64=_b64(_xlsx_bytes(tmp_path)),
        external_id="u-xlsx",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "ok"
    staged = _staged(vault)
    assert staged == [vault / UPLOAD_STAGING_RELDIR / f"{safe_stem('u-xlsx')}.xlsx"]
    fragments = _fragments(vault)
    assert len(fragments) == 1
    source = frontmatter.load(fragments[0])["source"]
    assert source["origin_key"] == f"00-Creek-Meta/adepthood/uploads/{staged[0].name}"


def test_resend_after_purge_restages_cleanly(tmp_path: Path) -> None:
    """RTBF then re-send: the purge erases the staged bytes; a re-send restores.

    End-to-end guard on the whole chain — the staged file only disappears if
    the upload ledger stamped ``source.origin_key`` AND the purge sweep's
    containment root covers the uploads dir.
    """
    vault = _vault(tmp_path)

    def _send() -> dict[str, Any]:
        """Send the same intimate bytes under a stable external id."""
        return upload_tool(
            vault_path=vault,
            filename="tender.txt",
            content_base64=_b64(b"a tender upload " + _SECRET),
            external_id="u-rtbf",
            timestamp=_TS,
            tier="intimate",
            privacy_tier_ceiling=TierCeiling.INTIMATE,
        )

    first = _send()
    assert first["status"] == "ok"
    staged = _upload_staged_path(vault, "u-rtbf", "tender.txt")
    assert staged.is_file()
    assert _leaking_files(vault, _SECRET) != []

    PurgeEngine(vault).purge_fragment(str(first["fragment_id"]))

    assert not staged.exists()
    assert _fragments(vault) == []
    assert _leaking_files(vault, _SECRET) == []

    resent = _send()

    assert resent["status"] == "ok"
    assert len(_fragments(vault)) == 1
    assert staged.is_file()


# ---- Review findings on PR #1230 ----


def test_resending_an_id_under_a_new_extension_is_refused(tmp_path: Path) -> None:
    """One ``external_id`` maps to one fragment, whatever the filename says.

    The staged path carries the extension (``route_to_ingestor`` dispatches on
    it), so the ledger key does too. A re-send under a different extension
    therefore lands on a path with no ledger record and reads as a *creation* —
    silently forking the id into a second live fragment and orphaning the
    first one's staged bytes, which a purge by the returned ``fragment_id``
    would then leave on disk. That is a hole in the RTBF coverage this tool
    exists to provide, so the fork is refused instead.
    """
    vault = _vault(tmp_path)

    first = upload_tool(
        vault_path=vault,
        filename="notes.txt",
        content_base64=_b64(b"The original document body."),
        external_id="u-swap",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )
    assert first["status"] == "ok"
    assert len(_staged(vault)) == 1

    second = upload_tool(
        vault_path=vault,
        filename="notes.md",
        content_base64=_b64(b"A different body under the same id."),
        external_id="u-swap",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert second["status"] == "refused"
    assert set(second) == {"status", "tool", "tier_ceiling", "reason"}
    assert ".txt" in second["reason"]
    # No second fragment, and no second staged file to orphan.
    assert len(_fragments(vault)) == 1
    assert len(_staged(vault)) == 1
    assert _leaking_files(vault, b"A different body under the same id.") == []


def test_an_intimate_upload_is_not_reachable_through_a_changed_extension(
    tmp_path: Path,
) -> None:
    """The #970 gate surveys the whole id, not just the incoming path.

    Ranking only the exact target path would admit an ``open`` caller at an id
    whose intimate fragment happens to be staged under another extension —
    the overwrite gate would see no ledger record there and read a creation.
    """
    vault = _vault(tmp_path)

    seeded = upload_tool(
        vault_path=vault,
        filename="tender.txt",
        content_base64=_b64(b"Something intimate."),
        external_id="u-mixed",
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.ALL,
    )
    assert seeded["status"] == "ok"

    attempt = upload_tool(
        vault_path=vault,
        filename="tender.md",
        content_base64=_b64(b"Benign replacement."),
        external_id="u-mixed",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert attempt["status"] == "refused"
    # The generic reason: it must not disclose the protected tier, and must not
    # fall through to the extension-conflict reason, which would.
    assert attempt["reason"] == GENERIC_ABOVE_CEILING_REASON
    assert len(_staged(vault)) == 1
    assert _leaking_files(vault, b"Benign replacement.") == []


def test_a_missing_vault_is_refused_without_writing_anything(tmp_path: Path) -> None:
    """The 'vault unavailable' refusal must not recreate the vault it denies.

    ``SourceLedger.load`` returns an EMPTY ledger rather than raising when the
    vault is gone, so the #970 gate read the call as a creation and let it
    through to staging, whose ``mkdir(parents=True)`` rebuilt the tree and
    wrote the caller's bytes into a vault the very same response called
    unavailable. Every other refusal in this module is side-effect-free.
    """
    missing = tmp_path / "gone"
    payload = b"Bytes that must never reach disk."

    result = upload_tool(
        vault_path=missing,
        filename="notes.txt",
        content_base64=_b64(payload),
        external_id="u-novault",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert set(result) == {"status", "tool", "tier_ceiling", "reason"}
    assert result["reason"] == "vault unavailable"
    # Nothing was created: not the vault, not the staging tree, not the audit log.
    assert not missing.exists()


def test_a_failed_ingest_leaves_no_unpurgeable_staged_bytes(tmp_path: Path) -> None:
    """A refusal must not strand a copy of the document no purge can reach.

    ``run_ingest`` reports ``written=0`` for content the binary heuristic
    drops, so the bytes are already staged when the refusal is decided. Left
    there they carry no ledger record, hence no ``source.origin_key`` on any
    fragment — and the RTBF sweep finds staged bytes only through that key.
    The document would survive every purge the vault offers.
    """
    vault = _vault(tmp_path)
    secret = b"unroutable\x00binary\x00secret"

    result = upload_tool(
        vault_path=vault,
        filename="mystery.bin",
        content_base64=_b64(secret),
        external_id="u-orphan",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "refused"
    assert _fragments(vault) == []
    assert _staged(vault) == []
    assert _leaking_files(vault, secret) == []


def test_a_corrected_extension_after_a_failed_ingest_is_accepted(
    tmp_path: Path,
) -> None:
    """The idempotency contract has to survive a failed first attempt.

    A caller whose first upload was refused must be able to retry the same
    ``external_id`` with a working filename. If the failed attempt's staged
    file lingered, the extension-conflict gate would refuse the retry and
    advise purging a fragment that was never created — so the only remedy
    would be abandoning the id, quietly breaking "the same id updates in
    place" for exactly the callers who need the retry.
    """
    vault = _vault(tmp_path)

    refused = upload_tool(
        vault_path=vault,
        filename="notes.bin",
        content_base64=_b64(b"unroutable\x00bytes"),
        external_id="u-retry",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )
    assert refused["status"] == "refused"

    retried = upload_tool(
        vault_path=vault,
        filename="notes.txt",
        content_base64=_b64(b"The same document, correctly typed."),
        external_id="u-retry",
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert retried["status"] == "ok"
    assert retried["action"] == "created"
    assert len(_fragments(vault)) == 1
    assert len(_staged(vault)) == 1


# --------------------------------------------------------------------------
# Multi-sheet workbooks (#1305)
#
# A `.xlsx` routes to `SpreadsheetIngestor`, which emits one fragment per
# non-empty sheet. Until #1305 all N shared one fragment id, so the ledgered
# upload path wrote exactly one fragment and reported the other N-1 as
# `unchanged` — the loss wearing a success message.
#
# The fix gives each sheet its own ledger key, which is what these tests are
# really about: the upload tool resolves an `external_id` back to its
# fragment through that key, and a staged path that now maps to N records
# instead of 0-or-1 is load-bearing for the #970 tier gate, for
# `_discard_unreferenced`, and for the audit log.
# --------------------------------------------------------------------------

_SHEET_NAMES = ("Budget", "Notes", "Q3")


def _multi_sheet_xlsx_bytes(tmp_path: Path, marker: str = "june") -> bytes:
    """Build a genuine three-sheet ``.xlsx`` and return its bytes.

    Args:
        tmp_path: Directory to build the workbook in.
        marker: Cell value written into every sheet, so a caller can produce
            byte-different content for the same sheet layout.

    Returns:
        The workbook's bytes.
    """
    path = tmp_path / f"multi-{marker}.xlsx"
    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = _SHEET_NAMES[0]
    for index, name in enumerate(_SHEET_NAMES):
        sheet = first if index == 0 else workbook.create_sheet(name)
        sheet["A1"] = "label"
        sheet["B1"] = "value"
        sheet["A2"] = f"{name}-{marker}"
        sheet["B2"] = index
    workbook.save(path)
    return path.read_bytes()


def _origin_keys(vault: Path) -> list[str]:
    """Return every fragment's ``source.origin_key``, sorted.

    Args:
        vault: Vault root.

    Returns:
        The sorted origin keys, with a missing key rendered as ``"<none>"``
        so an absent key fails an equality assertion loudly.
    """
    keys = []
    for path in _fragments(vault):
        source = frontmatter.load(path).metadata.get("source", {})
        key = source.get("origin_key") if isinstance(source, dict) else None
        keys.append(str(key) if key else "<none>")
    return sorted(keys)


def test_multi_sheet_upload_writes_one_fragment_per_sheet(tmp_path: Path) -> None:
    """An uploaded three-sheet workbook lands three fragments, not one.

    Measured at HEAD before the fix: ``created=1 unchanged=2``, one file on
    disk holding the FIRST sheet (``Budget``) — ``_write_model`` takes the
    lock, finds the id already written and returns the existing path, so
    writes 2..N are silent no-ops. The issue body's claim that the LAST
    sheet survives is backwards.
    """
    vault = _vault(tmp_path / "vault")

    result = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-multi",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "ok"
    assert result["action"] == "created"
    fragments = _fragments(vault)
    assert len(fragments) == 3, [p.name for p in fragments]
    # The heading names the STAGED file, not the caller's filename: only the
    # extension of a caller-supplied name is ever trusted, and only to route.
    staged_name = f"{safe_stem('u-multi')}.xlsx"
    headings = sorted(
        line
        for path in fragments
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.startswith("# ")
    )
    assert headings == [f"# {staged_name} — {name}" for name in sorted(_SHEET_NAMES)]


def test_multi_sheet_upload_reports_every_sheet_in_the_audit_log(
    tmp_path: Path,
) -> None:
    """``affected_fragment_ids`` must hold all three ids, each exactly once.

    At HEAD the tool resolves the staged path to at most ONE ledger record,
    so the response's ``fragment_id`` names a single sheet and the audit
    entry under-reports what the call touched. An audit record that names
    one of three fragments is a false record, and it is the record an RTBF
    request is answered from.
    """
    vault = _vault(tmp_path / "vault")

    result = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-multi-audit",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    affected = result["affected_fragment_ids"]
    assert len(affected) == 3
    assert len(set(affected)) == 3
    assert result["fragment_id"] == affected[0]
    on_disk = {frontmatter.load(p).metadata["id"] for p in _fragments(vault)}
    assert set(affected) == on_disk
    entry = next(e for e in _audit(vault) if e.get("affected_fragment_ids"))
    assert sorted(entry["affected_fragment_ids"]) == sorted(affected)


def test_multi_sheet_upload_resend_is_a_true_no_op(tmp_path: Path) -> None:
    """Re-sending identical bytes reports ``unchanged`` and adds no fragment.

    ``updated`` here would be the failure this issue is about: post-fix but
    pre-ledger-key, the run overwrote sheet 1 with sheet 2 and then sheet 3
    and reported ``updated=2``, which reads like work was done.
    """
    vault = _vault(tmp_path / "vault")
    payload = _b64(_multi_sheet_xlsx_bytes(tmp_path))

    def _send() -> dict[str, Any]:
        """Send the identical workbook under a stable external id."""
        return upload_tool(
            vault_path=vault,
            filename="book.xlsx",
            content_base64=payload,
            external_id="u-multi-resend",
            tier="open",
            timestamp=_TS,
            privacy_tier_ceiling=TierCeiling.OPEN,
        )

    first = _send()
    keys_after_first = _origin_keys(vault)
    second = _send()

    assert first["action"] == "created"
    assert second["action"] == "unchanged"
    assert len(_fragments(vault)) == 3
    assert len(_staged(vault)) == 1
    # The origin key is the RTBF purge key; a key that moves on re-ingest
    # silently drops the old fragment out of every purge the vault offers.
    assert _origin_keys(vault) == keys_after_first
    assert sorted(second["affected_fragment_ids"]) == sorted(
        first["affected_fragment_ids"]
    )


def test_multi_sheet_origin_keys_name_each_sheet(tmp_path: Path) -> None:
    """Each sheet fragment carries its own ``…/book.xlsx#<sheet>`` key."""
    vault = _vault(tmp_path / "vault")

    upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-multi-keys",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    staged_name = f"{safe_stem('u-multi-keys')}.xlsx"
    base = f"{UPLOAD_STAGING_RELDIR.as_posix()}/{staged_name}"
    assert _origin_keys(vault) == sorted(f"{base}#{name}" for name in _SHEET_NAMES)


def test_multi_sheet_resend_is_still_refused_above_ceiling(tmp_path: Path) -> None:
    """The #970 gate must still see an intimate workbook it cannot read.

    GREEN at HEAD (one collapsed fragment resolves fine), and it is the
    guard that the per-sheet ledger key does not defeat the gate. Resolving
    the staged path to a single record returns ``None`` once the ledger
    holds only ``…#<sheet>`` keys, which makes ``existing`` ``None`` and
    ADMITS a caller at ``ceiling=open`` to overwrite intimate content.
    Read this as a ratchet test, not as the issue's red test.
    """
    vault = _vault(tmp_path / "vault")
    first = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path, marker="tender")),
        external_id="u-multi-tier",
        timestamp=_TS,
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.INTIMATE,
    )
    assert first["status"] == "ok"
    assert first["fragment_id"] is not None
    assert len(first["affected_fragment_ids"]) == 3

    refused = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path, marker="benign")),
        external_id="u-multi-tier",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert refused["status"] == "refused"
    assert set(refused) == _REFUSAL_KEYS
    assert refused["reason"] == GENERIC_ABOVE_CEILING_REASON
    assert len(_fragments(vault)) == 3


def test_multi_sheet_extension_conflict_is_still_refused(tmp_path: Path) -> None:
    """A re-send of the same id under a new extension must still be refused.

    ``_refuse_extension_conflict`` asks whether the OTHER staged paths for
    this id resolve to a fragment. With per-sheet keys the bare staged path
    resolves to nothing, so a single-record lookup would answer "no
    conflict" and fork the id onto two live fragments, orphaning the
    workbook's staged bytes.
    """
    vault = _vault(tmp_path / "vault")
    upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-multi-ext",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    conflicted = upload_tool(
        vault_path=vault,
        filename="book.txt",
        content_base64=_b64(b"a plain text re-type of the same document"),
        external_id="u-multi-ext",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert conflicted["status"] == "refused"
    assert ".xlsx" in conflicted["reason"]
    assert len(_fragments(vault)) == 3
    assert len(_staged(vault)) == 1


def test_failed_resend_does_not_discard_a_referenced_workbook(
    tmp_path: Path,
) -> None:
    """``_discard_unreferenced`` must not unlink a workbook three sheets cite.

    Keyed on "does the ledger know this staged path", which with per-sheet
    keys is no longer a single-record question. Answering it with a
    single-record lookup deletes the staged bytes that three live fragments
    still reach through ``origin_key`` — the RTBF sweep's only route to
    them — on the first re-send whose ingest errors.
    """
    vault = _vault(tmp_path / "vault")
    upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-multi-fail",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )
    assert len(_staged(vault)) == 1

    failed = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path, marker="second")),
        external_id="u-multi-fail",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
        run=_failing_runner,
    )

    assert failed["status"] == "refused"
    assert len(_staged(vault)) == 1, "a referenced workbook must survive a failure"
    assert len(_fragments(vault)) == 3


def test_single_sheet_upload_keeps_a_bare_origin_key(tmp_path: Path) -> None:
    """A one-sheet workbook must be untouched by the sub-unit scheme.

    The single-sheet carve-out is the migration answer: every CSV and
    single-sheet XLSX already in an operator vault keeps its exact id, its
    exact filename and its exact ``origin_key``, so re-ingesting an
    untouched file after this change is still a no-op rather than a
    duplicate. This pins the key half of that.
    """
    vault = _vault(tmp_path / "vault")

    upload_tool(
        vault_path=vault,
        filename="budget.xlsx",
        content_base64=_b64(_xlsx_bytes(tmp_path)),
        external_id="u-single",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    staged_name = f"{safe_stem('u-single')}.xlsx"
    assert _origin_keys(vault) == [f"{UPLOAD_STAGING_RELDIR.as_posix()}/{staged_name}"]
    assert len(_fragments(vault)) == 1


def test_multi_sheet_fragment_id_is_stable_under_sheet_order(tmp_path: Path) -> None:
    """The response's singular ``fragment_id`` must not depend on sheet order.

    ``fragment_id`` reduces N sheets to one id, and callers persist it as
    the handle they later purge by. Reducing by *ledger-key order* makes
    that choice a property of the workbook; reducing by dict-insertion
    order makes it a property of the order the JSONL happened to be
    written and re-read in, which is not a promise the ledger makes.

    Sheet names are deliberately anti-alphabetical so the two orders
    disagree: insertion gives ``Zeta`` first, key order gives ``Alpha``.
    """
    vault = _vault(tmp_path / "vault")
    names = ("Zeta", "Alpha", "Mid")
    path = tmp_path / "ordered.xlsx"
    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = names[0]
    for index, name in enumerate(names):
        sheet = first if index == 0 else workbook.create_sheet(name)
        sheet["A1"] = "label"
        sheet["A2"] = name
    workbook.save(path)

    result = upload_tool(
        vault_path=vault,
        filename="ordered.xlsx",
        content_base64=_b64(path.read_bytes()),
        external_id="u-order",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    staged_name = f"{safe_stem('u-order')}.xlsx"
    base = f"{UPLOAD_STAGING_RELDIR.as_posix()}/{staged_name}"
    assert _origin_keys(vault) == sorted(f"{base}#{name}" for name in names)
    # Alphabetically first key, not workbook-first sheet.
    alpha_id = frontmatter.load(
        next(
            p
            for p in _fragments(vault)
            if frontmatter.load(p).metadata["source"]["origin_key"] == f"{base}#Alpha"
        )
    ).metadata["id"]
    assert result["fragment_id"] == alpha_id
    assert result["affected_fragment_ids"][0] == alpha_id


def test_purging_one_sheet_does_not_open_the_tier_gate_on_its_siblings(
    tmp_path: Path,
) -> None:
    """Purge one sheet, then re-upload at ``open``: still refused.

    The two-step privacy bypass this guards. Purging ONE sheet fragment of
    an N-sheet intimate upload deletes the shared staged workbook — that is
    the RTBF contract, the source bytes go — while the other N-1 sheet
    fragments stay live, intimate, and still ledgered.

    A #970 overwrite gate that surveys the STAGING DIRECTORY then sees
    nothing, concludes the ``external_id`` owns nothing, and admits a
    caller at ``ceiling=open`` as a *creation*. The re-upload restages to
    the same deterministic path, ``write_fragment_idempotent`` matches the
    surviving siblings by their intact ledger keys, and overwrites intimate
    content with the low-ceiling caller's bytes. Two individually
    legitimate operations, one tier bypass.

    A ledger record outlives the file it points at, which is why the survey
    reads the ledger.
    """
    vault = _vault(tmp_path / "vault")
    first = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path, marker="tender")),
        external_id="u-purge-gate",
        timestamp=_TS,
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.INTIMATE,
    )
    assert first["status"] == "ok"
    victim, *survivors = sorted(first["affected_fragment_ids"])
    assert len(survivors) == 2

    purge = PurgeEngine(vault).purge_fragment(victim)

    assert purge.journal_staged_removed == 1
    assert _staged(vault) == [], "the shared workbook is gone, by design"
    assert len(_fragments(vault)) == 2, "the sibling sheets are still live"

    refused = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path, marker="benign")),
        external_id="u-purge-gate",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert refused["status"] == "refused"
    assert set(refused) == _REFUSAL_KEYS
    assert refused["reason"] == GENERIC_ABOVE_CEILING_REASON
    # Nothing restaged, and the surviving intimate fragments are untouched.
    assert _staged(vault) == []
    assert len(_fragments(vault)) == 2
    assert {frontmatter.load(p).metadata["id"] for p in _fragments(vault)} == set(
        survivors
    )
    assert _leaking_files(vault, b"benign") == []


def test_extension_conflict_survives_a_purge_of_the_staged_bytes(
    tmp_path: Path,
) -> None:
    """A ledgered fragment whose staged bytes are gone is still a conflict.

    Same root cause as the tier-gate case: the conflict survey must read
    the ledger, not a directory listing. Otherwise purging one sheet lets
    the same ``external_id`` be re-sent under a different extension, forking
    it onto a second live fragment while the first's siblings remain.
    """
    vault = _vault(tmp_path / "vault")
    first = upload_tool(
        vault_path=vault,
        filename="book.xlsx",
        content_base64=_b64(_multi_sheet_xlsx_bytes(tmp_path)),
        external_id="u-purge-ext",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )
    PurgeEngine(vault).purge_fragment(sorted(first["affected_fragment_ids"])[0])
    assert _staged(vault) == []

    conflicted = upload_tool(
        vault_path=vault,
        filename="book.txt",
        content_base64=_b64(b"a plain text re-type of the same document"),
        external_id="u-purge-ext",
        tier="open",
        timestamp=_TS,
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert conflicted["status"] == "refused"
    assert ".xlsx" in conflicted["reason"]
    assert _staged(vault) == []
    assert len(_fragments(vault)) == 2


def test_a_tombed_ledger_record_still_closes_the_tier_gate(tmp_path: Path) -> None:
    """A soft-tombed record must not read as "this id owns nothing".

    Tombing is unreachable for uploads today — ``tomb_missing_units``
    requires both a ``TOMBING_SOURCES`` source type and a *directory*
    input, and ``creek.upload`` always passes a single file — so this
    seeds the record directly rather than pretending to reach it.

    The branch is kept because tombing is a soft delete: the fragment
    still exists under ``10-Liminal/Orphaned/``. Surveying ``live_keys``
    would hand a caller at ``ceiling=open`` a clean "creation" over
    content that is merely filed elsewhere, and the day uploads become
    tombable that would be a silent tier bypass rather than a new bug.
    """
    from creek.ingest.ledger import SourceLedger

    vault = _vault(tmp_path / "vault")
    first = upload_tool(
        vault_path=vault,
        filename="tender.txt",
        content_base64=_b64(b"a tender upload " + _SECRET),
        external_id="u-tombed",
        timestamp=_TS,
        tier="intimate",
        privacy_tier_ceiling=TierCeiling.INTIMATE,
    )
    assert first["status"] == "ok"

    ledger = SourceLedger.load(vault, source="upload")
    key = next(iter(ledger.live_keys()))
    record = ledger.get(key)
    assert record is not None
    ledger.record(key, record.fragment_id, record.content_hash, tombed=True)
    assert SourceLedger.load(vault, source="upload").live_keys() == set()

    refused = upload_tool(
        vault_path=vault,
        filename="tender.txt",
        content_base64=_b64(b"benign overwrite attempt"),
        external_id="u-tombed",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert refused["status"] == "refused"
    assert refused["reason"] == GENERIC_ABOVE_CEILING_REASON
    assert _leaking_files(vault, _SECRET) != [], "the intimate content survives"


def test_a_same_named_key_outside_the_staging_dir_does_not_block_an_upload(
    tmp_path: Path,
) -> None:
    """The survey is scoped to the staging dir, so a stem collision elsewhere
    cannot produce a phantom refusal.

    The upload ledger is shared machinery: ``derive_source_key`` falls back
    to a bare filename for an out-of-vault source, and a borrowed ledger
    could hold a key under some other directory. Matching on the filename
    alone would let such a key masquerade as this ``external_id``'s own
    content and refuse a legitimate first upload — a gate that fails
    *closed* on unrelated state is still a gate that is wrong.
    """
    from creek.ingest.ledger import SourceLedger

    vault = _vault(tmp_path / "vault")
    stem = safe_stem("u-scoped")
    decoy_ledger = SourceLedger.load(vault, source="upload")
    decoy_ledger.record(f"01-Fragments/{stem}.txt", "frag-decoy", "deadbeef")
    decoy_ledger.record(f"{stem}.txt", "frag-decoy-bare", "deadbeef")

    result = upload_tool(
        vault_path=vault,
        filename="notes.txt",
        content_base64=_b64(b"a perfectly ordinary first upload"),
        external_id="u-scoped",
        timestamp=_TS,
        tier="open",
        privacy_tier_ceiling=TierCeiling.OPEN,
    )

    assert result["status"] == "ok"
    assert result["action"] == "created"
    assert len(_fragments(vault)) == 1
    assert "frag-decoy" not in result["affected_fragment_ids"]
