"""Tests for ``creek.ingest.refresh`` — the FEAT-031 ``authored_at`` backfill.

Exercises the per-format dispatch table, the idempotency contract,
and the counter accounting in :class:`RefreshDatesResult`.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import TYPE_CHECKING

import frontmatter
import pytest

from creek.ingest.refresh import RefreshDatesResult, refresh_authored_dates
from creek.models import Fragment, FragmentSource, SourcePlatform

if TYPE_CHECKING:
    from pathlib import Path

# ---- Helpers -----------------------------------------------------------


@pytest.fixture()
def vault(tmp_path: Path) -> Path:
    """Return a vault root with a minimal ``01-Fragments/Notes/`` tree."""
    (tmp_path / "00-Creek-Meta").mkdir()
    (tmp_path / "01-Fragments" / "Notes").mkdir(parents=True)
    return tmp_path


def _write_fragment_file(
    vault: Path,
    name: str,
    *,
    original_file: str,
    authored_at: datetime | None = None,
    body: str = "Body.",
    platform: SourcePlatform = SourcePlatform.MARKDOWN,
) -> Path:
    """Materialise a Fragment to disk with the requested metadata.

    Returns the file path so each test can read it back and assert
    on the rewritten frontmatter / preserved body.
    """
    fragment = Fragment(
        id="frag-" + name[:12].ljust(12, "0"),
        title=name,
        source=FragmentSource(
            platform=platform,
            original_file=original_file,
        ),
        authored_at=authored_at,
    )
    data = fragment.model_dump(mode="json")
    post = frontmatter.Post(content=body, **data)
    path = vault / "01-Fragments" / "Notes" / f"{name}.md"
    path.write_text(frontmatter.dumps(post), encoding="utf-8")
    return path


# ---- Top-level behaviour ----------------------------------------------


class TestRefreshAuthoredDates:
    """End-to-end behaviour of the backfill orchestrator."""

    def test_returns_empty_result_when_no_fragments_dir(self, tmp_path: Path) -> None:
        """A vault root without ``01-Fragments/`` is a no-op (no errors)."""
        result = refresh_authored_dates(tmp_path)
        assert result == RefreshDatesResult()

    def test_updates_fragment_with_extractable_source(
        self, vault: Path, tmp_path: Path
    ) -> None:
        """Markdown frontmatter ``published:`` lands on ``authored_at``."""
        source = tmp_path / "note.md"
        source.write_text(
            "---\npublished: 2024-03-15\n---\n# Hi\n",
            encoding="utf-8",
        )
        frag_path = _write_fragment_file(vault, "alpha", original_file=str(source))

        result = refresh_authored_dates(vault)

        assert result.updated == 1
        post = frontmatter.load(str(frag_path))
        assert post.metadata["authored_at"].startswith("2024-03-15")
        # Body is preserved.
        assert post.content.strip() == "Body."

    def test_already_set_fragment_is_skipped(self, vault: Path, tmp_path: Path) -> None:
        """A fragment with ``authored_at`` already populated is left untouched."""
        source = tmp_path / "note.md"
        source.write_text(
            "---\npublished: 2024-03-15\n---\nIrrelevant\n",
            encoding="utf-8",
        )
        existing = datetime(2020, 1, 1, tzinfo=UTC)
        frag_path = _write_fragment_file(
            vault, "alpha", original_file=str(source), authored_at=existing
        )
        before = frag_path.read_bytes()

        result = refresh_authored_dates(vault)

        assert result.already_set == 1
        assert result.updated == 0
        # The file's bytes did not change — true idempotency.
        assert frag_path.read_bytes() == before

    def test_missing_source_is_counted(self, vault: Path) -> None:
        """A fragment whose ``original_file`` does not exist is counted as missing."""
        _write_fragment_file(
            vault, "gone", original_file="/nonexistent/path/somewhere.md"
        )

        result = refresh_authored_dates(vault)

        assert result.missing_source == 1
        assert result.updated == 0

    def test_empty_original_file_is_missing(self, vault: Path) -> None:
        """An empty ``original_file`` string counts as missing source."""
        _write_fragment_file(vault, "noop", original_file="")

        result = refresh_authored_dates(vault)

        assert result.missing_source == 1

    def test_unsupported_extension_is_counted(
        self, vault: Path, tmp_path: Path
    ) -> None:
        """JSON export (chat platform) is reported as unsupported, not missing."""
        chat_export = tmp_path / "claude.json"
        chat_export.write_text("[]", encoding="utf-8")
        _write_fragment_file(
            vault,
            "chat",
            original_file=str(chat_export),
            platform=SourcePlatform.CLAUDE,
        )

        result = refresh_authored_dates(vault)

        assert result.unsupported == 1
        assert result.missing_source == 0

    def test_no_extractable_date_is_counted(self, vault: Path, tmp_path: Path) -> None:
        """A source file with no extractable date increments ``no_date_found``."""
        source = tmp_path / "note.md"
        source.write_text("# Just a title\n\nNo frontmatter.\n", encoding="utf-8")
        _write_fragment_file(vault, "dateless", original_file=str(source))

        result = refresh_authored_dates(vault)

        assert result.no_date_found == 1
        assert result.updated == 0

    def test_txt_source_yields_no_date_found(self, vault: Path, tmp_path: Path) -> None:
        """TXT files carry no embedded date metadata → no_date_found."""
        source = tmp_path / "note.txt"
        source.write_text("Plain text.", encoding="utf-8")
        _write_fragment_file(vault, "plain", original_file=str(source))

        result = refresh_authored_dates(vault)

        assert result.no_date_found == 1

    def test_docx_creation_date_is_used(self, vault: Path, tmp_path: Path) -> None:
        """A real DOCX with core ``created`` populates ``authored_at`` on refresh.

        Picks DOCX over PDF for the per-format integration test because
        the test environment ships ``python-docx`` for both reading and
        writing — pdfminer.six can read but not write a PDF with the
        precise ``/CreationDate`` shape this test requires.
        """
        from docx import Document as DocxDocument

        docx_path = tmp_path / "doc.docx"
        doc = DocxDocument()
        doc.add_paragraph("Body.")
        doc.core_properties.created = datetime(2024, 3, 15, 8, 30, tzinfo=UTC)
        doc.save(docx_path)
        _write_fragment_file(vault, "report", original_file=str(docx_path))

        result = refresh_authored_dates(vault)

        assert result.updated == 1

    def test_scanned_total_matches_fragment_count(
        self, vault: Path, tmp_path: Path
    ) -> None:
        """``scanned`` counts every loadable fragment under 01-Fragments/.

        Sum of every disposition counter equals ``scanned`` — the
        contract a downstream summary relies on.
        """
        source_with_date = tmp_path / "with_date.md"
        source_with_date.write_text(
            "---\npublished: 2024-03-15\n---\nx\n", encoding="utf-8"
        )
        source_without_date = tmp_path / "no_date.md"
        source_without_date.write_text("# x\n", encoding="utf-8")
        _write_fragment_file(vault, "a", original_file=str(source_with_date))
        _write_fragment_file(vault, "b", original_file=str(source_without_date))
        _write_fragment_file(vault, "c", original_file="/missing/path.md")

        result = refresh_authored_dates(vault)

        assert result.scanned == 3
        assert (
            result.updated
            + result.already_set
            + result.no_date_found
            + result.missing_source
            + result.unsupported
            == result.scanned
        )

    def test_non_fragment_markdown_is_ignored(
        self, vault: Path, tmp_path: Path
    ) -> None:
        """Plain markdown notes coexisting with fragments are silently skipped."""
        # A non-fragment note that ``try_load_fragment`` returns ``None`` for.
        non_fragment = vault / "01-Fragments" / "Notes" / "stray.md"
        non_fragment.write_text(
            "---\ntype: not-a-fragment\n---\nPlain note.\n",
            encoding="utf-8",
        )

        result = refresh_authored_dates(vault)

        # The fragment-roundup logic decrements ``scanned`` for
        # non-fragments so accounting stays honest.
        assert result.scanned == 0
        assert result.updated == 0


class TestRefreshPerFormatHandlers:
    """Per-format file-handler unit tests.

    Each handler reads a source file off disk and routes to the
    appropriate per-format extractor. These tests target the file-I/O
    branches that the orchestration tests above only cover incidentally.
    """

    def test_html_handler_extracts_meta_published_time(self, tmp_path: Path) -> None:
        """``_html_authored_at_from_file`` reads OG metadata."""
        from creek.ingest.refresh import _html_authored_at_from_file

        html_path = tmp_path / "post.html"
        html_path.write_text(
            "<html><head>"
            '<meta property="article:published_time" '
            'content="2024-03-15T08:30:00Z">'
            "</head><body>x</body></html>",
            encoding="utf-8",
        )
        result = _html_authored_at_from_file(html_path)
        assert result is not None
        assert result == datetime(2024, 3, 15, 8, 30, tzinfo=UTC)

    def test_html_handler_returns_none_on_io_error(self, tmp_path: Path) -> None:
        """Missing source file → ``None`` from the HTML handler."""
        from creek.ingest.refresh import _html_authored_at_from_file

        result = _html_authored_at_from_file(tmp_path / "missing.html")
        assert result is None

    def test_html_handler_returns_none_when_no_metadata(self, tmp_path: Path) -> None:
        """HTML without date metadata → ``None``."""
        from creek.ingest.refresh import _html_authored_at_from_file

        html_path = tmp_path / "post.html"
        html_path.write_text("<html><body>plain</body></html>", encoding="utf-8")
        assert _html_authored_at_from_file(html_path) is None

    def test_pdf_handler_returns_none_on_io_error(self, tmp_path: Path) -> None:
        """Missing PDF source file → ``None``."""
        from creek.ingest.refresh import _pdf_authored_at_from_file

        assert _pdf_authored_at_from_file(tmp_path / "missing.pdf") is None

    def test_docx_handler_extracts_core_created(self, tmp_path: Path) -> None:
        """A real DOCX with ``created`` core property is parsed correctly."""
        from docx import Document as DocxDocument

        from creek.ingest.refresh import _docx_authored_at_from_file

        docx_path = tmp_path / "doc.docx"
        doc = DocxDocument()
        doc.add_paragraph("x")
        doc.core_properties.created = datetime(2024, 3, 15, 8, 30, tzinfo=UTC)
        doc.save(docx_path)
        result = _docx_authored_at_from_file(docx_path)
        assert result is not None
        assert result.date() == datetime(2024, 3, 15).date()

    def test_docx_handler_returns_none_on_io_error(self, tmp_path: Path) -> None:
        """Missing DOCX source file → ``None``."""
        from creek.ingest.refresh import _docx_authored_at_from_file

        assert _docx_authored_at_from_file(tmp_path / "missing.docx") is None

    def test_rtf_handler_extracts_creatim(self, tmp_path: Path) -> None:
        """RTF ``\\creatim`` control word is parsed correctly."""
        from creek.ingest.refresh import _rtf_authored_at_from_file

        rtf_path = tmp_path / "doc.rtf"
        rtf_path.write_text(
            "{\\rtf1\\ansi"
            "{\\info{\\creatim\\yr2024\\mo3\\dy15\\hr8\\min30\\sec0}}"
            "Body.}",
            encoding="ascii",
        )
        result = _rtf_authored_at_from_file(rtf_path)
        assert result is not None
        assert result == datetime(2024, 3, 15, 8, 30, tzinfo=UTC)

    def test_rtf_handler_returns_none_on_io_error(self, tmp_path: Path) -> None:
        """Missing RTF source file → ``None``."""
        from creek.ingest.refresh import _rtf_authored_at_from_file

        assert _rtf_authored_at_from_file(tmp_path / "missing.rtf") is None

    def test_markdown_handler_extracts_published(self, tmp_path: Path) -> None:
        """Markdown ``published:`` is parsed via the shared extraction chain."""
        from creek.ingest.refresh import _markdown_authored_at_from_file

        md_path = tmp_path / "note.md"
        md_path.write_text("---\npublished: 2024-03-15\n---\nx\n", encoding="utf-8")
        result = _markdown_authored_at_from_file(md_path)
        assert result is not None
        assert result.date() == datetime(2024, 3, 15).date()

    def test_txt_handler_always_returns_none(self, tmp_path: Path) -> None:
        """TXT handler returns ``None`` by definition (FEAT-031)."""
        from creek.ingest.refresh import _txt_authored_at_from_file

        # Path doesn't need to exist — handler short-circuits.
        assert _txt_authored_at_from_file(tmp_path / "x.txt") is None

    def test_resolve_unsupported_extension_returns_none(self, tmp_path: Path) -> None:
        """Dispatcher returns ``None`` for unsupported extensions."""
        from creek.ingest.refresh import _resolve_authored_at_for_source

        path = tmp_path / "data.bin"
        path.write_bytes(b"\x00\x01\x02")
        assert _resolve_authored_at_for_source(path) is None

    def test_resolve_swallows_handler_exceptions(self, tmp_path: Path) -> None:
        """Handler exceptions are logged and swallowed → ``None``.

        Drive this through a real DOCX dispatch: handing the DOCX
        handler a non-DOCX file makes ``python-docx`` raise, which
        the dispatcher must convert to ``None`` rather than propagate.
        """
        from creek.ingest.refresh import _resolve_authored_at_for_source

        path = tmp_path / "broken.docx"
        path.write_bytes(b"not a real docx")
        assert _resolve_authored_at_for_source(path) is None


class TestRefreshAuthoredDatesIdempotency:
    """A second backfill pass touches nothing — pins the FEAT-031 contract."""

    def test_second_pass_changes_no_bytes(self, vault: Path, tmp_path: Path) -> None:
        """File bytes are byte-identical across two consecutive runs."""
        source = tmp_path / "note.md"
        source.write_text(
            "---\npublished: 2024-03-15\n---\nx\n",
            encoding="utf-8",
        )
        frag_path = _write_fragment_file(vault, "alpha", original_file=str(source))

        refresh_authored_dates(vault)
        snapshot1 = frag_path.read_bytes()
        refresh_authored_dates(vault)
        snapshot2 = frag_path.read_bytes()

        assert snapshot1 == snapshot2

    def test_second_pass_reports_zero_updates(
        self, vault: Path, tmp_path: Path
    ) -> None:
        """Counter accounting reflects the no-op nature of the second pass."""
        source = tmp_path / "note.md"
        source.write_text(
            "---\npublished: 2024-03-15\n---\nx\n",
            encoding="utf-8",
        )
        _write_fragment_file(vault, "alpha", original_file=str(source))

        refresh_authored_dates(vault)
        second_result = refresh_authored_dates(vault)

        assert second_result.updated == 0
        assert second_result.already_set == 1
